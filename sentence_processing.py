# sentence_processing.py

from flask import g, current_app
from flask_login import current_user
from rapidfuzz import fuzz
import re
from docx import Document
from spacy_manager import SpacyModel
from models import db, Paragraph, KeyWord, Report, HeadSentence, BodySentence, TailSentence, HeadSentenceGroup, BodySentenceGroup, TailSentenceGroup
from logger import logger
from collections import defaultdict



# Функция для обработки строки ключевых слов, разделенных запятой
def process_keywords(key_word_input: str) -> list:
    """Обрабатываем строку ключевых слов, разделенных запятой, 
    и возвращаем список"""
    
    key_words = []
    for word in key_word_input.split(','):
        stripped_word = word.strip()
        if stripped_word:
            key_words.append(stripped_word)
    return key_words


def check_existing_keywords(key_words):
    """
    Проверяет, какие ключевые слова уже существуют у текущего профиля и возвращает сообщение об ошибке
    в случае дублирования, или None, если ключевые слова уникальны.
    
    Args:
        key_words (list): Список новых ключевых слов.

    Returns:
        string or None: Сообщение об ошибке, если дублирующиеся ключевые слова найдены, или None, если ключевые слова уникальны.
    """
    
    profile_key_words = KeyWord.find_by_profile(g.current_profile.id)

    # Создаем словарь для хранения результатов
    existing_keywords = {}

    # Проходим по всем ключевым словам пользователя
    for kw in profile_key_words:
        key_word_lower = kw.key_word.lower()
        if key_word_lower in [kw.lower() for kw in key_words]:
            # Проверяем, связаны ли ключевые слова с отчетами
            if kw.key_word_reports:
                for report in kw.key_word_reports:
                    report_name = report.report_name
                    if report_name not in existing_keywords:
                        existing_keywords[report_name] = []
                    existing_keywords[report_name].append(kw.key_word)
            else:
                # Если слово не связано с отчетом, оно является глобальным
                if "global" not in existing_keywords:
                    existing_keywords["global"] = []
                existing_keywords["global"].append(kw.key_word)

    # Если дублирующиеся ключевые слова найдены, формируем сообщение
    if existing_keywords:
        existing_keywords_message = []
        for key, words in existing_keywords.items():
            words_list = ", ".join(words)
            if key == "global":
                existing_keywords_message.append(f"global: {words_list}")
            else:
                existing_keywords_message.append(f"{key}: {words_list}")

        message = "The following keywords already exist:\n" + "\n".join(existing_keywords_message)
        return message

    return None


# Функция для извлечения абзацев и предложений из документа Word.
# Используется в new_report_creation.py
def extract_paragraphs_and_sentences(file_path):
    """
    Извлекает заголовки и предложения из документа Word.

    Функция обрабатывает документ .docx, определяя абзацы, начинающиеся с жирного текста, как заголовки разделов. 
    Остальные строки считаются содержимым этих разделов. Предложения внутри абзацев могут разделяться символом '!!', 
    формируя альтернативные варианты одного и того же предложения.

    Args:
        file_path (str): Путь к файлу Word.

    Returns:
        list: Список словарей, представляющих абзацы документа. 
              Каждый словарь содержит:
              - 'title' (str): Заголовок абзаца (если есть).
              - 'sentences' (list): Список предложений (или списков альтернативных предложений).
              - 'visible' (bool): Видимость абзаца (по умолчанию True).
              - 'bold' (bool): Является ли заголовок жирным (по умолчанию False).
              - 'is_title' (bool): Помечен ли абзац как заголовок (по умолчанию False).
              - 'type_paragraph' (str): Тип абзаца (по умолчанию "text").
    """
    document = Document(file_path)
    # Проверяю не пустой ли документ
    if not document.paragraphs:
        return []
    
    paragraphs_from_file = []
    current_paragraph = None

    for para in document.paragraphs:
        # Check if paragraph has bold text indicating it's a new section
        if para.runs and para.runs[0].bold:
            # Handle previous paragraph before moving to the next one
            if current_paragraph:
                paragraphs_from_file.append(current_paragraph)
            
            # Create a new paragraph entry
            current_paragraph = {
                'title': para.text.strip(),
                'sentences': []
            }
        else:
            # Append sentences to the current paragraph
            if current_paragraph:
                sentences = para.text.strip().split('!!')
                # If there are multiple parts, store them as a nested list
                if len(sentences) > 1:
                    current_paragraph['sentences'].append([s.strip() for s in sentences])
                else:
                    current_paragraph['sentences'].append(sentences[0].strip())

    # Append the last paragraph
    if current_paragraph:
        paragraphs_from_file.append(current_paragraph)

    return paragraphs_from_file


# Предварительная очистка. Проверяю на наличие предложения. 
# Убираю только повторяющиеся знаки и лишние пробелы. 
# Исключение - многоточие, обрабатывается отдельно. 
# Добавляю пробелы после знаков препинания, если их нет
def preprocess_sentence(text):
    """
    Performs a rough cleanup of the text to prepare it for further processing.
    Args:
        text (str): The input sentence or text.
    Returns:
        str: The preprocessed text.
    """
    # Проверяем, содержит ли текст хотя бы одну букву или цифру
    if not re.search(r'[a-zA-Zа-яА-ЯёЁ0-9]', text):
        return ""

   # Обрабатываем повторяющиеся знаки, кроме точки
    text = re.sub(r'([,!?:;"\'\(\)])\1+', r'\1', text)  # Повторяющиеся знаки → один

    # Обрабатываем многоточия отдельно
    text = re.sub(r'\.\.\.\.+', '...', text)  # Четыре и более точек → многоточие
    text = re.sub(r'(?<!\.)\.\.(?!\.)', '.', text)  # Две точки → одна
    
    # Добавляю пробелы после знаков препинания, если их нет
    text = re.sub(r'([.!?,;:])(?!\s)', r'\1 ', text)
    
    # Убираем лишние пробелы
    text = re.sub(r'\s+', ' ', text)  # Заменяем любые пробельные символы на один пробел

    return text


# это функция очистки текста перед сравнением 
# используется в working_with_reports.py дважды
def clean_text_with_keywords(sentence, key_words, except_words=None):
    """ Функция очистки текста от лишних пробелов, знаков припенания, 
    цифр, слов исключений и ключевых слов с приведением всех слов 
    предложения к нижнему регистру """
    
    
    
    
    # Приводим текст к строчным буквам
    sentence = str(sentence).lower()
    
    # Удаляем ключевые слова
    if key_words:
        for word in key_words:
            sentence = re.sub(rf'(?<!\w){re.escape(word)}(?!\w)', '', sentence, flags=re.IGNORECASE)
    
    # Удаляем все кроме букв цифр и робелов
    sentence = re.sub(r"[^\w\s]", "", sentence)
    # Убираем цифры
    sentence = re.sub(r"\d+", "", sentence)
    # Удаляем лишние пробелы
    sentence = re.sub(r"\s+", " ", sentence).strip()
    # Убираем дополнительные слова
    if except_words:
        for word in except_words:
            sentence = re.sub(rf"(?<!\w){re.escape(word)}(?!\w)", "", sentence, flags=re.IGNORECASE)
    
    return sentence


# это функция для окончательной очистки предложения перед сохранением в базу данных
# в working_with_reports.py.
# Она не очищает двойные знаки, лишние пробелы и не проверяется текст на наличие
# Должна применяться после функции preprocess_sentence
def clean_and_normalize_text(text):
    """
    Cleans the input text by handling punctuation, spaces, and formatting issues.
    Args:
        text (str): The text to be cleaned.
    Returns:
        str: The cleaned and normalized text.
    """
    
    # Исключения для слов, которые должны начинаться с заглавной буквы
    exeptions_after_punctuation =current_app.config["PROFILE_SETTINGS"]["EXCEPTIONS_AFTER_PUNCTUATION"]

    # Убираем пронумерованные элементы с точкой или скобкой в начале строки
    # Это в тему именно тут, так как ниже я обрабатываю скобки
    text = re.sub(r'^\s*\d+[\.\)]\s*', '', text)
    
    # Убираем лишние пробелы
    text = re.sub(r'\s+([.,!?;:])', r'\1', text)  # Пробел перед знаками препинания
    text = re.sub(r'\(\s+', r'(', text)  # Пробел после открывающей скобки
    text = re.sub(r'\s+\)', r')', text)  # Пробел перед закрывающей скобкой

    # Убираем пробелы в начале и конце строки. 
    # Это в тему именну тут, так как выше мы добавляем пробелы.
    text = text.strip()

    # Проверяем, начинается ли предложение с заглавной буквы
    if text and text[0].islower():
        text = text[0].upper() + text[1:]
        
    # Проверяем слова после двоеточий и запятых
    def process_after_punctuation(match):
        """
        Обрабатывает слово после знака препинания.
        Если слово не в исключениях, переводит его в нижний регистр.
        """
        punctuation = match.group(1)
        word = match.group(2)
        if word in exeptions_after_punctuation:
            return f"{punctuation} {word}"  # Оставляем слово как есть
        return f"{punctuation} {word.lower()}"  # Приводим к нижнему регистру
    
    # Регулярное выражение для поиска знаков препинания и последующего слова
    text = re.sub(r'([,:]) (\w+)', process_after_punctuation, text)
        
    if not re.search(r'[.!?]$', text):
        text += '.'

    return text


# разделяет текс по предложениям, функция используется в working_with_reports.py
def split_sentences_if_needed(text):
    """
    Splits a sentence into multiple sentences using SpaCy tokenizer.
    Args:
        text (str): The input sentence text.
    Returns:
        tuple: (list of valid sentences, list of excluded sentences).
    """
    logger.info(f"(функция split_sentences_if_needed) -----------------------------------------")
    logger.info(f"(функция split_sentences_if_needed) 🚀 Начато разбиение текста на предложения")
    language = current_app.config.get("PROFILE_SETTINGS", {}).get("APP_LANGUAGE", "ru")
    logger.info(f"(функция split_sentences_if_needed) получен текст: {text}")
    # Загружаем модель SpaCy для текущего языка
    try:
        nlp = SpacyModel.get_instance(language)
    except ValueError as e:
        logger.error(f"(функция split_sentences_if_needed) ❌ Ошибка загрузки модели SpaCy: {e}")
        return [], []
    
    doc = nlp(text)
    sentences = [sent.text for sent in doc.sents]
    for sentence in sentences:
        if not re.search(r'[a-zA-Zа-яА-ЯёЁ0-9]', sentence):
            logger.info(f"(функция split_sentences_if_needed) ⚠️ Пропускаю предложение в котором нет букв или цифр: ({sentence})")
            sentences.remove(sentence)
    logger.info(f"(функция split_sentences_if_needed) После разделения предложений мы получили -  ({sentences}) ")
    if len(sentences) > 1:
        # Если найдено более одного предложения, добавляем в исключения
        return [], sentences # Splitted sentences
    return sentences, [] # Unsplitted sentences



def sort_key_words_group(unsorted_key_words_group):
    """
    Сортировка групп ключевых слов по первой букве первого ключевого слова в группе.
    Работает для всех вариантов структуры: с отчетами, с индексами и без них.
    """
    key_words_group_with_first_letter = []

    for group_data in unsorted_key_words_group:
        # В зависимости от структуры данных выбираем ключевые слова
        key_words = group_data.get("key_words", [])

        if key_words:
            # Проверяем, как представлены ключевые слова — как строки или объекты с полем "word"
            if isinstance(key_words[0], dict):
                # Если ключевые слова — это словари, берем поле "word"
                first_letter = key_words[0]["word"][0].lower() if key_words[0]["word"] else ""
            else:
                # Если ключевые слова — это строки
                first_letter = key_words[0][0].lower() if key_words[0] else ""
        else:
            first_letter = ""  # Если ключевых слов нет, используем пустую строку

        # Добавляем первую букву и саму группу данных для сортировки
        key_words_group_with_first_letter.append((first_letter, group_data))

    # Сортируем список по первой букве
    sorted_key_words_group = sorted(key_words_group_with_first_letter, key=lambda x: x[0])

    # Извлекаем только отсортированные группы (без первой буквы)
    return [group_data for _, group_data in sorted_key_words_group]


def group_keywords(keywords, with_index=False, with_report=False):
    """
    Создаем список сгруппированных слов или список словарей 
    с ключами group_index и key_words в зависимости от полученных 
    на входе аргументов, а также список ключевых слов с отчетами,
    если with_report=True.

    Args:
        keywords (list): Список объектов класса KeyWordGroup.
        with_index (boolean): Флаг для определения результата с или без добавления индекса.
        with_report (boolean): Флаг для группировки ключевых слов по отчетам.
    
    Returns:
    with_index=True:
        list of dict: Список словарей с ключами group_index и key_words (включает id и слово).
    with_index=False:
        list of lists: Список списков с отсортированными по group_index ключевыми словами (включает id и слово).
    with_report=True:
        list of dict: Список словарей, где ключи report_id, report_name, group_index и key_words (включает id и слово).
    """
    
    # Группируем ключевые слова по group_index, добавляя их id и слово
    grouped_keywords = defaultdict(list)
    for keyword in keywords:
        grouped_keywords[keyword.group_index].append({'id': keyword.id, 'word': keyword.key_word})

    if with_report:
        # Используем список для хранения данных по отчетам
        report_key_words = []
        seen_groups = set()  # Множество для хранения уникальных групп

        for keyword in keywords:
            group_index = keyword.group_index
            group_words = grouped_keywords[group_index]

            # Для каждого отчета, связанного с ключевым словом
            for report in keyword.key_word_reports:
                # Проверяем, не была ли группа уже добавлена для этого отчета
                if (report.id, group_index) not in seen_groups:
                    # Добавляем всю группу ключевых слов, если её еще не добавили
                    report_key_words.append({
                        'report_id': report.id,
                        'report_name': report.report_name,
                        'group_index': group_index,
                        'key_words': group_words  # Добавляем всю группу ключевых слов с их id
                    })
                    # Отмечаем, что группа с этим group_index уже добавлена для данного отчета
                    seen_groups.add((report.id, group_index))

        return report_key_words

    # Если with_index=True, возвращаем список словарей с group_index
    if with_index:
        grouped_keywords_with_index = []
        for group_index, words in grouped_keywords.items():
            grouped_keywords_with_index.append({'group_index': group_index, 'key_words': words})
        return grouped_keywords_with_index
    
    # Если with_index=False, возвращаем просто сгруппированные ключевые слова с их id
    return list(grouped_keywords.values())


# Сравниваю 2 предложения. Используется в working_with_report/save_modified_sentences. 
# Ищет совпадения с заданным порогом, также очищает текст от чисел и ключевых слов
def compare_sentences_by_paragraph(new_sentences, report_id):    
    """
    Compares new sentences with existing sentences in their respective paragraphs to determine uniqueness.
    """
    logger.info(f"(функция compare_sentences_by_paragraph) 🚀 Начато сравнение новых предложений с существующими в базе данных")
    logger.debug(f"(функция compare_sentences_by_paragraph) Получены новые предложения - ({new_sentences})")
    similarity_threshold_fuzz = int(current_app.config["PROFILE_SETTINGS"]["SIMILARITY_THRESHOLD_FUZZ"])
    except_words = current_app.config["PROFILE_SETTINGS"]["EXCEPT_WORDS"]
    logger.debug(f"(функция compare_sentences_by_paragraph) Порог схожести: {similarity_threshold_fuzz}")
    logger.info(f"(функция compare_sentences_by_paragraph) Исключаемые слова: {except_words}")
    
    existing_paragraphs = Paragraph.query.filter_by(report_id=report_id).all()
    key_words_obj = KeyWord.get_keywords_for_report(g.current_profile.id, report_id)
    key_words = [keyword.key_word for keyword in key_words_obj]
    logger.debug(f"(функция compare_sentences_by_paragraph) Получено {len(key_words)} ключевых слов.")
    
    duplicates = []
    unique_sentences = []
    errors_count = 0
    # Итерируем по новым предложениям
    for new_sentence in new_sentences:
        new_paragraph_id = int(new_sentence.get("paragraph_id"))
        new_text = new_sentence.get("text")
        new_sentence_type = new_sentence.get("sentence_type")
        new_sentence_head_sentence_id = new_sentence.get("head_sentence_id") or None
        
        if not new_paragraph_id or not new_text.strip():
            errors_count += 1
            logger.warning(f"(функция compare_sentences_by_paragraph) ⚠️ Пропускаю предложение ({new_sentence}) с пустым текстом или пустым id параграфа")
            continue  # Пропускаем некорректные данные
        
        # Находим соответствующий параграф
        paragraph = next((p for p in existing_paragraphs if p.id == new_paragraph_id), None)
        # Находим соответствующую группу для предложения
        related_group_id = None
        sentence_group_class = None
        
        if new_sentence_type == "body":
            head_sentence = HeadSentence.query.get(new_sentence_head_sentence_id)
            if not head_sentence:
                logger.warning(f"(функция compare_sentences_by_paragraph) ⚠️ Главное предложение с id={new_sentence_head_sentence_id} не найдено. Пропускаю предложение ({new_sentence})")
                errors_count += 1
                continue
            related_group_id = head_sentence.body_sentence_group_id or None
            existing_sentences = BodySentenceGroup.get_group_sentences(related_group_id) or []
            logger.debug(f"(функция compare_sentences_by_paragraph) {existing_sentences}")
            existing_sentences.append({"id": head_sentence.id, "sentence": head_sentence.sentence})
            
        else:
            if not paragraph:
                logger.warning(f"(функция compare_sentences_by_paragraph) ⚠️ Параграф с id={new_paragraph_id} не найден. Пропускаю предложение ({new_sentence})")
                errors_count += 1
                continue
            related_group_id = paragraph.tail_sentence_group_id or None
            if not related_group_id:
                logger.warning(f"(функция compare_sentences_by_paragraph) Группа хвостовых предложений не найдена. Добавляю в уникальные")
                unique_sentences.append(new_sentence)
                continue
            existing_sentences = TailSentenceGroup.get_group_sentences(related_group_id)
            
            logger.debug(f"(функция compare_sentences_by_paragraph) {existing_sentences}")
        
        # Очищаем существующие предложения
        cleaned_existing = []
        for sent in existing_sentences:
            cleaned_item = {
                "id": sent.get("id"),
                "original_text": sent["sentence"],
                "cleaned_text": clean_text_with_keywords(sent.get("sentence"), key_words, except_words)
            }
            cleaned_existing.append(cleaned_item)
            
        # Очищаем новое предложение
        cleaned_new_text = clean_text_with_keywords(new_text, key_words, except_words)
        # Проверяем на схожесть с существующими предложениями
        is_duplicate = False
        for existing in cleaned_existing:
            similarity_rapidfuzz = fuzz.ratio(cleaned_new_text, existing["cleaned_text"])
            if similarity_rapidfuzz >= similarity_threshold_fuzz:
                duplicates.append({
                    "new_sentence": new_sentence,
                   
                    "new_sentence_paragraph": new_sentence.get("paragraph_id"),
                    "matched_with": {
                                "id": existing["id"],
                                "text": existing["original_text"]
                            },
                    "similarity_rapidfuzz": similarity_rapidfuzz
                })
                is_duplicate = True
                logger.debug(f"(функция compare_sentences_by_paragraph) 🔄 Найдено дублирующее предложение ({new_sentence}) с существующим ({existing['original_text']}) с похожестью {similarity_rapidfuzz}")
                break

        if not is_duplicate:
            unique_sentences.append(new_sentence)

    return {"duplicates": duplicates, "unique": unique_sentences, "errors_count": errors_count}


# Функция проверяет предложение на уникальность и добавляет его 
# в результат, если оно уникально. Работает со списком по ссылке, 
# ничего не возвращает, просто меняет предоставленные список и множество
def _add_if_unique(raw_text, key_words, except_words, cleaned_list, result_set, threshold):
    """
    Проверяет, является ли предложение уникальным, и добавляет его в результат.

    Args:
        raw_text (str): Оригинальный текст.
        key_words (list): Ключевые слова для очистки.
        except_words (list): Слова-исключения.
        cleaned_list (list): Уже очищенные тексты.
        result_set (set): Уникальные предложения (результат).
        threshold (int): Порог схожести.
    """
    cleaned = clean_text_with_keywords(raw_text, key_words, except_words)
    for existing in cleaned_list:
        if fuzz.ratio(cleaned, existing) >= threshold:
            return
    cleaned_list.append(cleaned)
    result_set.add(raw_text)
   


# Функция для поиска существующих аналогичных предложений того же типа в базе данных 
# использую в models.py. Ищет 100% совпадения
def find_similar_exist_sentence(sentence_text, sentence_type, report_type_id):
    """
    Ищет похожие предложения в базе данных.
    """
    user_id = current_user.id
    tags = None
    logger.info(f"(функция find_similar_exist_sentence)(тип предложения: '{sentence_type}') 🚀 Начат поиск существующих аналогичных предложений в базе данных")
    
    # Получаем все предложения того же типа и с такими же базовыми параметрами из базы данных
    if sentence_type == "head":
        similar_type_sentences = HeadSentence.query.filter_by(tags=tags, report_type_id=report_type_id, user_id = user_id).all()
    elif sentence_type == "body":
        similar_type_sentences = BodySentence.query.filter_by(tags=tags, report_type_id=report_type_id, user_id = user_id).all()
    elif sentence_type == "tail":
        similar_type_sentences = TailSentence.query.filter_by(tags=tags, report_type_id=report_type_id, user_id = user_id).all()
    else:
        raise ValueError(f"Invalid sentence type: {sentence_type}")
    
    # Сравниваем входное предложение с каждым из существующих
    for exist_sentence in similar_type_sentences:
        if exist_sentence.sentence == sentence_text:
            logger.info(f"(функция find_similar_exist_sentence) 🧩 Найдено совпадение с предложением '{exist_sentence.sentence}' в базе данных. Возвращаю предложение")
            return exist_sentence
    logger.info(f"(функция find_similar_exist_sentence) Совпадений не найдено. Возвращаю None")
    return None
    
      
      
# Функция для оценки уникальности индексов в главных предложениях параграфа. Номера индексов не должны повторяться
def check_head_sentence_indexes(paragraph_id):
    """
    Проверяет уникальность индексов главных предложений в параграфе.
    """
    head_sentences_groupe, _ = Paragraph.get_paragraph_groups(paragraph_id)
    
    head_sentences = HeadSentenceGroup.get_sentences(head_sentences_groupe)
    indexes = [sent.index for sent in head_sentences]
    duplicates = [index for index in indexes if indexes.count(index) > 1]
    return duplicates

        
      
        