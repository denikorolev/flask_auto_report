# file_processing.py

from flask import g
from flask_login import current_user
from werkzeug.datastructures import FileStorage
from docx import Document
import os
from unidecode import unidecode
from datetime import datetime
import tempfile
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from config import get_config, Config
from models import db, FileMetadata


# Проверка допустимости расширения загружаемого файла
def allowed_file(file_name, file_type):
    if file_type == "doc":
        return '.' in file_name and file_name.rsplit('.', 1)[1].lower() in {'docx', "doc"}
    elif file_type == "tab":
        return '.' in file_name and file_name.rsplit('.', 1)[1].lower() in {"xlsx"}
    elif file_type == "jpg":
        return '.' in file_name and file_name.rsplit('.', 1)[1].lower() in {"jpg", "jpeg", "png"}
    else:
        return False

# Преобразование пути к файлу в объект файла
def create_filestorage_from_path(file_path):
    """
    Преобразует файл с диска в объект FileStorage, который Flask использует для загруженных файлов.
    
    :param file_path: Путь к файлу на диске.
    :return: Объект FileStorage.
    """
    try:
        # Открываем файл в режиме чтения
        file = open(file_path, "rb")
        # Создаем объект FileStorage, передавая файл и его имя
        file_storage = FileStorage(file, filename=os.path.basename(file_path))
        return file_storage
    except Exception as e:
        print(f"Error while creating FileStorage: {e}")
        return None

def sanitize_filename(filename):
    """
    Удаляет пробелы, тире и выполняет транслитерацию имени файла.
    
    :param filename: Исходное имя файла.
    :return: Обработанное имя файла.
    """
    # Заменяем пробелы и тире на подчеркивания
    sanitized_name = filename.replace(" ", "_").replace("-", "_")
    # Выполняем транслитерацию (убираем специальные символы)
    return unidecode(sanitized_name)


def sync_profile_files(profile_id):
    """
    Синхронизирует файлы профиля: проверяет записи в базе данных и соответствующие файлы на диске.
    Удаляет записи в базе данных, если файла нет на диске, и удаляет файлы, если нет записи в базе данных.
    
    :param profile_id: ID профиля, для которого проводится синхронизация.
    """
    try:
        # Получаем папку профиля
        upload_folder = Config.get_user_upload_folder()

        # 1. Проверяем наличие всех файлов из базы данных
        files_in_db = FileMetadata.query.filter_by(profile_id=profile_id).all()
        
        for file_meta in files_in_db:
            file_path = file_meta.file_path

            if not os.path.exists(file_path):
                print(f"File {file_path} not found, deleting record from database...")
                db.session.delete(file_meta)
                db.session.commit()

        # 2. Проверяем наличие всех файлов в папке профиля
        for root, dirs, files in os.walk(upload_folder):
            for file in files:
                file_path = os.path.join(root, file)

                # Проверяем, есть ли этот файл в базе данных
                file_record = FileMetadata.query.filter_by(file_path=file_path).first()

                if not file_record:
                    print(f"File {file_path} not found in database, deleting file from filesystem...")
                    os.remove(file_path)

        return "Synchronization completed successfully."

    except Exception as e:
        print(f"Error during file synchronization: {e}")
        return f"Error during file synchronization: {e}"


def file_uploader(file, file_type, folder_name, file_name=None, file_description=None):
    """
    Uploads a file to the server in the folder corresponding to the user's profile and specified folder name,
    and saves metadata to the database.
    
    :param file: файл, который загружается или путь к файлу
    :param file_type: тип файла (например, "doc" или "jpg")
    :param folder_name: имя папки, в которую будет загружен файл (например, "templates" или "word_reports")
    :param file_name: имя файла (если не передано, используется имя папки)
    :param file_description: описание файла для записи в базу данных
    
    :return: сообщение об успехе или ошибке и путь к файлу.
    """
    # Проверяем, является ли file строкой, если да, то это путь к файлу
    if isinstance(file, str):
        file = create_filestorage_from_path(file)
        if not file:
            return "Failed to convert file path to FileStorage", None

    # Проверка, есть ли у файла атрибут filename
    if not hasattr(file, "filename") or file.filename == "":
        return "The file name is empty", None

    # Проверка допустимого расширения файла
    if not allowed_file(file.filename, file_type):
        return "The file name or extension is not allowed", None
    
    # Получаем расширение файла
    name, ext = os.path.splitext(file.filename)
    
    # Если имя файла не передано, используем имя папки в качестве имени файла
    if not file_name:
        file_name = folder_name
    sanitized_name = sanitize_filename(file_name)
    # Формируем строку с текущей датой и временем
    date_str = datetime.now().strftime("%d_%m_%y")
    time_str = datetime.now().strftime("%H_%M")
    
    # Создаем имя файла с датой и временем
    filename_with_date_time = f"{sanitized_name}_{date_str}_{time_str}{ext}"
    
    # Получаем путь к папке пользователя с учетом текущего профиля
    try:
        user_folder = get_config().get_user_upload_folder()
    except Exception as e:
        return f"Failed to get user folder: {e}", None
    
    # Путь к папке, соответствующей folder_name и текущей дате
    target_folder = os.path.join(user_folder, folder_name, f"{folder_name}_{date_str}")
    
    # Создаем папку, если она не существует
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
    
    # Полный путь к файлу
    file_path = os.path.join(target_folder, filename_with_date_time)
    
    try:
        # Сохраняем файл по указанному пути
        file.save(file_path)

        # Создаем запись о файле в базе данных
        file_description = file_description or folder_name
        profile_id = g.current_profile.id
        try:
            FileMetadata.create(
                profile_id=profile_id,
                file_name=filename_with_date_time,
                file_path=file_path,
                file_type=file_type,
                file_description=file_description
            )
        except Exception as e:
            return f"Can't add data to file metadata: {e}", None
        
        return "The file was uploaded successfully and metadata saved", file_path
    
    except Exception as e:
        return f"The file wasn't uploaded due to the following error: {e}", None


def save_to_word(text, name, subtype, report_type, birthdate, reportnumber, scanParam, side=""):
    # Убираем лишние пробелы и проверяем на пустые строки
    name = name.strip() or "NoName"
    subtype = subtype.strip() or "NoSubtype"
    report_type = report_type.strip() or "NoReportType"

    try:
        date_str = datetime.now().strftime("%d_%m_%y")
        modified_date_str = date_str.replace("_", ".")
        # Преобразование строки в объект datetime
        trans_birthdate = datetime.strptime(birthdate, "%Y-%m-%d")
        # Форматирование даты
        modified_birthdate = trans_birthdate.strftime("%d.%m.%y")
    except:
        modified_birthdate = "unknown"
    try:
        profile_id = g.current_profile.id
        
        signatura_metadata = FileMetadata.get_file_by_description(profile_id, "signatura")
        template_metadata = FileMetadata.get_file_by_description(profile_id, "word_template")
        
        if not template_metadata:
            raise Exception("Word template not found in the database.")
        
        signatura_path = signatura_metadata.file_path
        template_path = template_metadata.file_path

        document = Document(template_path)
        
        def set_font_size(doc, size):
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size)
                    
                    
        # Insert the required text at the beginning of the document
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Протокол МРТ-Исследования № {reportnumber}")
        run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add patient information
        p = document.add_paragraph()
        p.add_run("ФИО пациента: ").bold = True
        p.add_run(name)
        p.paragraph_format.space_before = Pt(0)  
        p.paragraph_format.space_after = Pt(0)

        p = document.add_paragraph()
        p.add_run("Дата рождения: ").bold = True
        p.add_run(modified_birthdate)
        p.add_run(" г.р.")
        p.paragraph_format.space_before = Pt(0)  
        p.paragraph_format.space_after = Pt(0)
        
        p = document.add_paragraph()
        p.add_run("Вид исследования: ").bold = True
        p.add_run(report_type)
        
        # Добавляем сторону, если она указана
        if side == "right":
            p.add_run(" правого")
        elif side == "left":
            p.add_run(" левого")
            
        p.add_run(" ")
        p.add_run(subtype)
        p.paragraph_format.space_before = Pt(0)  
        p.paragraph_format.space_after = Pt(0)
        
        p = document.add_paragraph()
        p.add_run("Техника сканирования: ").bold = True
        p.add_run(scanParam)
        p.paragraph_format.space_before = Pt(0)  
        p.paragraph_format.space_after = Pt(0)
        
        document.add_paragraph(text)

        # Create a table with one row and three columns
        table = document.add_table(rows=1, cols=3)
        table.autofit = False
        
        # Set column widths
        widths = [Inches(3), Inches(2), Inches(2)]
        for col_idx, width in enumerate(widths):
            for cell in table.columns[col_idx].cells:
                cell.width = width
                
        # First cell: Doctor's name
        cell1 = table.cell(0, 0)
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell1.paragraphs[0]
        run = paragraph.add_run("Врач-рентгенолог: Королёв Д.Г. ")
        run.bold = True
        run.font.size = Pt(12)

        # Second cell: Image
        cell2 = table.cell(0, 1)
        paragraph = cell2.paragraphs[0]
        run = paragraph.add_run()
        if signatura_path:
            run.add_picture(signatura_path, width=Pt(70))

        cell3 = table.cell(0, 2)
        cell3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell3.paragraphs[0]
        run = paragraph.add_run("Дата:  ")
        run.bold = True
        run.font.size = Pt(12)
        run = paragraph.add_run(modified_date_str)
        run.font.size = Pt(12)

        # Set font size for the rest of the document
        set_font_size(document, 12)
        

        # Генерация имени файла
        filename = f"{name}_{report_type}_{subtype}"
        

        # Сохраняем документ во временный файл
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            temp_file_path = tmp_file.name
            document.save(temp_file_path)  # Сохраняем файл в эту временную директорию

        
        folder_name = "reports_word"
          
        result, saved_file_path = file_uploader(temp_file_path, "doc", folder_name, filename)
        
        # Удаляем временный файл после загрузки
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
            
        if "successfully" in result:
            print(f"File saved at: {saved_file_path}")
            return saved_file_path
        else:
            raise Exception(result)
    except Exception as e:
        print(f"Error in save_to_word: {e}")
        raise Exception(f"Error in save_to_word: {e}")



