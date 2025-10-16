# sentence_processing.py

from flask_security import current_user
from rapidfuzz import fuzz
import re
import json
from docx import Document
from sqlalchemy import select
from app.utils.spacy_manager import SpacyModel
from app.models.models import db, Paragraph, KeyWord, HeadSentence, BodySentence, TailSentence, HeadSentenceGroup, BodySentenceGroup, TailSentenceGroup, AppConfig, head_sentence_group_link
from app.utils.logger import logger
from collections import defaultdict



# –§—É–Ω–∫—Ü–∏—è –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ —Å—Ç—Ä–æ–∫–∏ –∫–ª—é—á–µ–≤—ã—Ö —Å–ª–æ–≤, —Ä–∞–∑–¥–µ–ª–µ–Ω–Ω—ã—Ö –∑–∞–ø—è—Ç–æ–π
def process_keywords(key_word_input: str) -> list:
    """–û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º —Å—Ç—Ä–æ–∫—É –∫–ª—é—á–µ–≤—ã—Ö —Å–ª–æ–≤, —Ä–∞–∑–¥–µ–ª–µ–Ω–Ω—ã—Ö –∑–∞–ø—è—Ç–æ–π, 
    –∏ –≤–æ–∑–≤—Ä–∞—â–∞–µ–º —Å–ø–∏—Å–æ–∫"""
    
    key_words = []
    for word in key_word_input.split(','):
        stripped_word = word.strip()
        if stripped_word:
            key_words.append(stripped_word)
    return key_words


def check_existing_keywords(key_words, profile_id):
    """
    –ü—Ä–æ–≤–µ—Ä—è–µ—Ç, –∫–∞–∫–∏–µ –∫–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞ —É–∂–µ —Å—É—â–µ—Å—Ç–≤—É—é—Ç —É —Ç–µ–∫—É—â–µ–≥–æ –ø—Ä–æ—Ñ–∏–ª—è –∏ –≤–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Å–æ–æ–±—â–µ–Ω–∏–µ –æ–± –æ—à–∏–±–∫–µ
    –≤ —Å–ª—É—á–∞–µ –¥—É–±–ª–∏—Ä–æ–≤–∞–Ω–∏—è, –∏–ª–∏ None, –µ—Å–ª–∏ –∫–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞ —É–Ω–∏–∫–∞–ª—å–Ω—ã.
    
    Args:
        key_words (list): –°–ø–∏—Å–æ–∫ –Ω–æ–≤—ã—Ö –∫–ª—é—á–µ–≤—ã—Ö —Å–ª–æ–≤.

    Returns:
        string or None: –°–æ–æ–±—â–µ–Ω–∏–µ –æ–± –æ—à–∏–±–∫–µ, –µ—Å–ª–∏ –¥—É–±–ª–∏—Ä—É—é—â–∏–µ—Å—è –∫–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞ –Ω–∞–π–¥–µ–Ω—ã, –∏–ª–∏ None, –µ—Å–ª–∏ –∫–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞ —É–Ω–∏–∫–∞–ª—å–Ω—ã.
    """
    profile_key_words = KeyWord.find_by_profile(profile_id)

    # –°–æ–∑–¥–∞–µ–º —Å–ª–æ–≤–∞—Ä—å –¥–ª—è —Ö—Ä–∞–Ω–µ–Ω–∏—è —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤
    existing_keywords = {}

    # –ü—Ä–æ—Ö–æ–¥–∏–º –ø–æ –≤—Å–µ–º –∫–ª—é—á–µ–≤—ã–º —Å–ª–æ–≤–∞–º –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è
    for kw in profile_key_words:
        key_word_lower = kw.key_word.lower()
        if key_word_lower in [kw.lower() for kw in key_words]:
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —Å–≤—è–∑–∞–Ω—ã –ª–∏ –∫–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞ —Å –æ—Ç—á–µ—Ç–∞–º–∏
            if kw.key_word_reports:
                for report in kw.key_word_reports:
                    report_name = report.report_name
                    if report_name not in existing_keywords:
                        existing_keywords[report_name] = []
                    existing_keywords[report_name].append(kw.key_word)
            else:
                # –ï—Å–ª–∏ —Å–ª–æ–≤–æ –Ω–µ —Å–≤—è–∑–∞–Ω–æ —Å –æ—Ç—á–µ—Ç–æ–º, –æ–Ω–æ —è–≤–ª—è–µ—Ç—Å—è –≥–ª–æ–±–∞–ª—å–Ω—ã–º
                if "global" not in existing_keywords:
                    existing_keywords["global"] = []
                existing_keywords["global"].append(kw.key_word)

    # –ï—Å–ª–∏ –¥—É–±–ª–∏—Ä—É—é—â–∏–µ—Å—è –∫–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞ –Ω–∞–π–¥–µ–Ω—ã, —Ñ–æ—Ä–º–∏—Ä—É–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ
    if existing_keywords:
        existing_keywords_message = []
        for key, words in existing_keywords.items():
            words_list = ", ".join(words)
            if key == "global":
                existing_keywords_message.append(f"global: {words_list}")
            else:
                existing_keywords_message.append(f"{key}: {words_list}")

        message = "The following keywords already exist:\n" + "\n".join(existing_keywords_message)
        return message

    return None


# –§—É–Ω–∫—Ü–∏—è –¥–ª—è –∏–∑–≤–ª–µ—á–µ–Ω–∏—è –∞–±–∑–∞—Ü–µ–≤ –∏ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π –∏–∑ –¥–æ–∫—É–º–µ–Ω—Ç–∞ Word.
# –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –≤ new_report_creation.py
def extract_paragraphs_and_sentences(file_path):
    """
    –ò–∑–≤–ª–µ–∫–∞–µ—Ç –∑–∞–≥–æ–ª–æ–≤–∫–∏ –∏ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è –∏–∑ –¥–æ–∫—É–º–µ–Ω—Ç–∞ Word.

    –§—É–Ω–∫—Ü–∏—è –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ—Ç –¥–æ–∫—É–º–µ–Ω—Ç .docx, –æ–ø—Ä–µ–¥–µ–ª—è—è –∞–±–∑–∞—Ü—ã, –Ω–∞—á–∏–Ω–∞—é—â–∏–µ—Å—è —Å –∂–∏—Ä–Ω–æ–≥–æ —Ç–µ–∫—Å—Ç–∞, –∫–∞–∫ –∑–∞–≥–æ–ª–æ–≤–∫–∏ —Ä–∞–∑–¥–µ–ª–æ–≤. 
    –û—Å—Ç–∞–ª—å–Ω—ã–µ —Å—Ç—Ä–æ–∫–∏ —Å—á–∏—Ç–∞—é—Ç—Å—è —Å–æ–¥–µ—Ä–∂–∏–º—ã–º —ç—Ç–∏—Ö —Ä–∞–∑–¥–µ–ª–æ–≤. –ü—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è –≤–Ω—É—Ç—Ä–∏ –∞–±–∑–∞—Ü–µ–≤ –º–æ–≥—É—Ç —Ä–∞–∑–¥–µ–ª—è—Ç—å—Å—è —Å–∏–º–≤–æ–ª–æ–º '!!', 
    —Ñ–æ—Ä–º–∏—Ä—É—è –∞–ª—å—Ç–µ—Ä–Ω–∞—Ç–∏–≤–Ω—ã–µ –≤–∞—Ä–∏–∞–Ω—Ç—ã –æ–¥–Ω–æ–≥–æ –∏ —Ç–æ–≥–æ –∂–µ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è.

    Args:
        file_path (str): –ü—É—Ç—å –∫ —Ñ–∞–π–ª—É Word.

    Returns:
        list: –°–ø–∏—Å–æ–∫ —Å–ª–æ–≤–∞—Ä–µ–π, –ø—Ä–µ–¥—Å—Ç–∞–≤–ª—è—é—â–∏—Ö –∞–±–∑–∞—Ü—ã –¥–æ–∫—É–º–µ–Ω—Ç–∞. 
              –ö–∞–∂–¥—ã–π —Å–ª–æ–≤–∞—Ä—å —Å–æ–¥–µ—Ä–∂–∏—Ç:
              - 'title' (str): –ó–∞–≥–æ–ª–æ–≤–æ–∫ –∞–±–∑–∞—Ü–∞ (–µ—Å–ª–∏ –µ—Å—Ç—å).
              - 'sentences' (list): –°–ø–∏—Å–æ–∫ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π (–∏–ª–∏ —Å–ø–∏—Å–∫–æ–≤ –∞–ª—å—Ç–µ—Ä–Ω–∞—Ç–∏–≤–Ω—ã—Ö –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π).
              - 'visible' (bool): –í–∏–¥–∏–º–æ—Å—Ç—å –∞–±–∑–∞—Ü–∞ (–ø–æ —É–º–æ–ª—á–∞–Ω–∏—é True).
              - 'bold' (bool): –Ø–≤–ª—è–µ—Ç—Å—è –ª–∏ –∑–∞–≥–æ–ª–æ–≤–æ–∫ –∂–∏—Ä–Ω—ã–º (–ø–æ —É–º–æ–ª—á–∞–Ω–∏—é False).
              - 'is_title' (bool): –ü–æ–º–µ—á–µ–Ω –ª–∏ –∞–±–∑–∞—Ü –∫–∞–∫ –∑–∞–≥–æ–ª–æ–≤–æ–∫ (–ø–æ —É–º–æ–ª—á–∞–Ω–∏—é False).
              - 'type_paragraph' (str): –¢–∏–ø –∞–±–∑–∞—Ü–∞ (–ø–æ —É–º–æ–ª—á–∞–Ω–∏—é "text").
    """
    document = Document(file_path)
    # –ü—Ä–æ–≤–µ—Ä—è—é –Ω–µ –ø—É—Å—Ç–æ–π –ª–∏ –¥–æ–∫—É–º–µ–Ω—Ç
    if not document.paragraphs:
        return []
    
    paragraphs_from_file = []
    current_paragraph = None

    for para in document.paragraphs:
        # Check if paragraph has bold text indicating it's a new section
        if para.runs and para.runs[0].bold:
            # Handle previous paragraph before moving to the next one
            if current_paragraph:
                paragraphs_from_file.append(current_paragraph)
            
            # Create a new paragraph entry
            current_paragraph = {
                'title': para.text.strip(),
                'sentences': []
            }
        else:
            # Append sentences to the current paragraph
            if current_paragraph:
                sentences = para.text.strip().split('!!')
                # If there are multiple parts, store them as a nested list
                if len(sentences) > 1:
                    current_paragraph['sentences'].append([s.strip() for s in sentences])
                else:
                    current_paragraph['sentences'].append(sentences[0].strip())

    # Append the last paragraph
    if current_paragraph:
        paragraphs_from_file.append(current_paragraph)

    return paragraphs_from_file


# –ü—Ä–µ–¥–≤–∞—Ä–∏—Ç–µ–ª—å–Ω–∞—è –æ—á–∏—Å—Ç–∫–∞. –ü—Ä–æ–≤–µ—Ä—è—é –Ω–∞ –Ω–∞–ª–∏—á–∏–µ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è. 
# –£–±–∏—Ä–∞—é —Ç–æ–ª—å–∫–æ –ø–æ–≤—Ç–æ—Ä—è—é—â–∏–µ—Å—è –∑–Ω–∞–∫–∏ –∏ –ª–∏—à–Ω–∏–µ –ø—Ä–æ–±–µ–ª—ã. 
# –ò—Å–∫–ª—é—á–µ–Ω–∏–µ - –º–Ω–æ–≥–æ—Ç–æ—á–∏–µ, –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ—Ç—Å—è –æ—Ç–¥–µ–ª—å–Ω–æ. 
# –î–æ–±–∞–≤–ª—è—é –ø—Ä–æ–±–µ–ª—ã –ø–æ—Å–ª–µ –∑–Ω–∞–∫–æ–≤ –ø—Ä–µ–ø–∏–Ω–∞–Ω–∏—è, –µ—Å–ª–∏ –∏—Ö –Ω–µ—Ç
def preprocess_sentence(text):
    """
    Performs a rough cleanup of the text to prepare it for further processing.
    Args:
        text (str): The input sentence or text.
    Returns:
        str: The preprocessed text.
    """
    # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —Å–æ–¥–µ—Ä–∂–∏—Ç –ª–∏ —Ç–µ–∫—Å—Ç —Ö–æ—Ç—è –±—ã –æ–¥–Ω—É –±—É–∫–≤—É –∏–ª–∏ —Ü–∏—Ñ—Ä—É
    if not re.search(r'[a-zA-Z–∞-—è–ê-–Ø—ë–Å0-9]', text):
        logger.info(f"(preprocess_sentence) ‚ö†Ô∏è –ü—Ä–æ–ø—É—Å–∫–∞—é –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ: {text}")
        return ""

    if len(text.strip()) < 3:
        logger.info(f"(preprocess_sentence) ‚ö†Ô∏è –ü—Ä–æ–ø—É—Å–∫–∞—é —Å–ª–∏—à–∫–æ–º –∫–æ—Ä–æ—Ç–∫–∏–π —Ç–µ–∫—Å—Ç: {text}")
        return ""

   # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º –ø–æ–≤—Ç–æ—Ä—è—é—â–∏–µ—Å—è –∑–Ω–∞–∫–∏, –∫—Ä–æ–º–µ —Ç–æ—á–∫–∏
    text = re.sub(r'([,!?:;"\'\(\)])\1+', r'\1', text)  # –ü–æ–≤—Ç–æ—Ä—è—é—â–∏–µ—Å—è –∑–Ω–∞–∫–∏ ‚Üí –æ–¥–∏–Ω

    # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º –º–Ω–æ–≥–æ—Ç–æ—á–∏—è –æ—Ç–¥–µ–ª—å–Ω–æ
    text = re.sub(r'\.\.\.\.+', '...', text)  # –ß–µ—Ç—ã—Ä–µ –∏ –±–æ–ª–µ–µ —Ç–æ—á–µ–∫ ‚Üí –º–Ω–æ–≥–æ—Ç–æ—á–∏–µ
    text = re.sub(r'(?<!\.)\.\.(?!\.)', '.', text)  # –î–≤–µ —Ç–æ—á–∫–∏ ‚Üí –æ–¥–Ω–∞
    
    # –î–æ–±–∞–≤–ª—è—é –ø—Ä–æ–±–µ–ª—ã –ø–æ—Å–ª–µ –∑–Ω–∞–∫–æ–≤ –ø—Ä–µ–ø–∏–Ω–∞–Ω–∏—è, –µ—Å–ª–∏ –∏—Ö –Ω–µ—Ç
    text = re.sub(r'([.!?,;:])(?!\s)', r'\1 ', text)
    
    # –£–±–∏—Ä–∞–µ–º –ª–∏—à–Ω–∏–µ –ø—Ä–æ–±–µ–ª—ã
    text = re.sub(r'\s+', ' ', text)  # –ó–∞–º–µ–Ω—è–µ–º –ª—é–±—ã–µ –ø—Ä–æ–±–µ–ª—å–Ω—ã–µ —Å–∏–º–≤–æ–ª—ã –Ω–∞ –æ–¥–∏–Ω –ø—Ä–æ–±–µ–ª

    return text


# —ç—Ç–æ —Ñ—É–Ω–∫—Ü–∏—è –æ—á–∏—Å—Ç–∫–∏ —Ç–µ–∫—Å—Ç–∞ –ø–µ—Ä–µ–¥ —Å—Ä–∞–≤–Ω–µ–Ω–∏–µ–º 
# –∏—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –≤ working_with_reports.py –¥–≤–∞–∂–¥—ã
def clean_text_with_keywords(sentence, key_words, except_words=None):
    """ –§—É–Ω–∫—Ü–∏—è –æ—á–∏—Å—Ç–∫–∏ —Ç–µ–∫—Å—Ç–∞ –æ—Ç –ª–∏—à–Ω–∏—Ö –ø—Ä–æ–±–µ–ª–æ–≤, –∑–Ω–∞–∫–æ–≤ –ø—Ä–∏–ø–µ–Ω–∞–Ω–∏—è, 
    —Ü–∏—Ñ—Ä, —Å–ª–æ–≤ –∏—Å–∫–ª—é—á–µ–Ω–∏–π –∏ –∫–ª—é—á–µ–≤—ã—Ö —Å–ª–æ–≤ —Å –ø—Ä–∏–≤–µ–¥–µ–Ω–∏–µ–º –≤—Å–µ—Ö —Å–ª–æ–≤ 
    –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è –∫ –Ω–∏–∂–Ω–µ–º—É —Ä–µ–≥–∏—Å—Ç—Ä—É """
    
    
    
    
    # –ü—Ä–∏–≤–æ–¥–∏–º —Ç–µ–∫—Å—Ç –∫ —Å—Ç—Ä–æ—á–Ω—ã–º –±—É–∫–≤–∞–º
    sentence = str(sentence).lower()
    
    # –£–¥–∞–ª—è–µ–º –∫–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞
    if key_words:
        for word in key_words:
            sentence = re.sub(rf'(?<!\w){re.escape(word)}(?!\w)', '', sentence, flags=re.IGNORECASE)
    
    # –£–¥–∞–ª—è–µ–º –≤—Å–µ –∫—Ä–æ–º–µ –±—É–∫–≤ —Ü–∏—Ñ—Ä –∏ —Ä–æ–±–µ–ª–æ–≤
    sentence = re.sub(r"[^\w\s]", "", sentence)
    # –£–±–∏—Ä–∞–µ–º —Ü–∏—Ñ—Ä—ã
    sentence = re.sub(r"\d+", "", sentence)
    # –£–¥–∞–ª—è–µ–º –ª–∏—à–Ω–∏–µ –ø—Ä–æ–±–µ–ª—ã
    sentence = re.sub(r"\s+", " ", sentence).strip()
    # –£–±–∏—Ä–∞–µ–º –¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–µ —Å–ª–æ–≤–∞
    if except_words:
        for word in except_words:
            sentence = re.sub(rf"(?<!\w){re.escape(word)}(?!\w)", "", sentence, flags=re.IGNORECASE)
    
    return sentence


# —ç—Ç–æ —Ñ—É–Ω–∫—Ü–∏—è –¥–ª—è –æ–∫–æ–Ω—á–∞—Ç–µ–ª—å–Ω–æ–π –æ—á–∏—Å—Ç–∫–∏ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è –ø–µ—Ä–µ–¥ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ–º –≤ –±–∞–∑—É –¥–∞–Ω–Ω—ã—Ö
# –≤ working_with_reports.py.
# –û–Ω–∞ –Ω–µ –æ—á–∏—â–∞–µ—Ç –¥–≤–æ–π–Ω—ã–µ –∑–Ω–∞–∫–∏, –ª–∏—à–Ω–∏–µ –ø—Ä–æ–±–µ–ª—ã –∏ –Ω–µ –ø—Ä–æ–≤–µ—Ä—è–µ—Ç—Å—è —Ç–µ–∫—Å—Ç –Ω–∞ –Ω–∞–ª–∏—á–∏–µ
# –î–æ–ª–∂–Ω–∞ –ø—Ä–∏–º–µ–Ω—è—Ç—å—Å—è –ø–æ—Å–ª–µ —Ñ—É–Ω–∫—Ü–∏–∏ preprocess_sentence
def clean_and_normalize_text(text, profile_id):
    """
    Cleans the input text by handling punctuation, spaces, and formatting issues.
    Args:
        text (str): The text to be cleaned.
    Returns:
        str: The cleaned and normalized text.
    """
    
    # –ò—Å–∫–ª—é—á–µ–Ω–∏—è –¥–ª—è —Å–ª–æ–≤, –∫–æ—Ç–æ—Ä—ã–µ –¥–æ–ª–∂–Ω—ã –Ω–∞—á–∏–Ω–∞—Ç—å—Å—è —Å –∑–∞–≥–ª–∞–≤–Ω–æ–π –±—É–∫–≤—ã
    exeptions_after_punctuation = AppConfig.get_setting(profile_id, "EXCEPTIONS_AFTER_PUNCTUATION", "").split(",")

    # –£–±–∏—Ä–∞–µ–º –ø—Ä–æ–Ω—É–º–µ—Ä–æ–≤–∞–Ω–Ω—ã–µ —ç–ª–µ–º–µ–Ω—Ç—ã —Å —Ç–æ—á–∫–æ–π –∏–ª–∏ —Å–∫–æ–±–∫–æ–π –≤ –Ω–∞—á–∞–ª–µ —Å—Ç—Ä–æ–∫–∏
    # –≠—Ç–æ –≤ —Ç–µ–º—É –∏–º–µ–Ω–Ω–æ —Ç—É—Ç, —Ç–∞–∫ –∫–∞–∫ –Ω–∏–∂–µ —è –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞—é —Å–∫–æ–±–∫–∏
    text = re.sub(r'^\s*\d+[\.\)]\s*', '', text)
    
    # –£–±–∏—Ä–∞–µ–º –ª–∏—à–Ω–∏–µ –ø—Ä–æ–±–µ–ª—ã
    text = re.sub(r'\s+([.,!?;:])', r'\1', text)  # –ü—Ä–æ–±–µ–ª –ø–µ—Ä–µ–¥ –∑–Ω–∞–∫–∞–º–∏ –ø—Ä–µ–ø–∏–Ω–∞–Ω–∏—è
    text = re.sub(r'\(\s+', r'(', text)  # –ü—Ä–æ–±–µ–ª –ø–æ—Å–ª–µ –æ—Ç–∫—Ä—ã–≤–∞—é—â–µ–π —Å–∫–æ–±–∫–∏
    text = re.sub(r'\s+\)', r')', text)  # –ü—Ä–æ–±–µ–ª –ø–µ—Ä–µ–¥ –∑–∞–∫—Ä—ã–≤–∞—é—â–µ–π —Å–∫–æ–±–∫–æ–π

    # –£–±–∏—Ä–∞–µ–º –ø—Ä–æ–±–µ–ª—ã –≤ –Ω–∞—á–∞–ª–µ –∏ –∫–æ–Ω—Ü–µ —Å—Ç—Ä–æ–∫–∏. 
    # –≠—Ç–æ –≤ —Ç–µ–º—É –∏–º–µ–Ω–Ω—É —Ç—É—Ç, —Ç–∞–∫ –∫–∞–∫ –≤—ã—à–µ –º—ã –¥–æ–±–∞–≤–ª—è–µ–º –ø—Ä–æ–±–µ–ª—ã.
    text = text.strip()

    # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –Ω–∞—á–∏–Ω–∞–µ—Ç—Å—è –ª–∏ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ —Å –∑–∞–≥–ª–∞–≤–Ω–æ–π –±—É–∫–≤—ã
    if text and text[0].islower():
        text = text[0].upper() + text[1:]
        
    # –ü—Ä–æ–≤–µ—Ä—è–µ–º —Å–ª–æ–≤–∞ –ø–æ—Å–ª–µ –¥–≤–æ–µ—Ç–æ—á–∏–π –∏ –∑–∞–ø—è—Ç—ã—Ö
    def process_after_punctuation(match):
        """
        –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ—Ç —Å–ª–æ–≤–æ –ø–æ—Å–ª–µ –∑–Ω–∞–∫–∞ –ø—Ä–µ–ø–∏–Ω–∞–Ω–∏—è.
        –ï—Å–ª–∏ —Å–ª–æ–≤–æ –Ω–µ –≤ –∏—Å–∫–ª—é—á–µ–Ω–∏—è—Ö, –ø–µ—Ä–µ–≤–æ–¥–∏—Ç –µ–≥–æ –≤ –Ω–∏–∂–Ω–∏–π —Ä–µ–≥–∏—Å—Ç—Ä.
        """
        punctuation = match.group(1)
        word = match.group(2)
        if word in exeptions_after_punctuation:
            return f"{punctuation} {word}"  # –û—Å—Ç–∞–≤–ª—è–µ–º —Å–ª–æ–≤–æ –∫–∞–∫ –µ—Å—Ç—å
        return f"{punctuation} {word.lower()}"  # –ü—Ä–∏–≤–æ–¥–∏–º –∫ –Ω–∏–∂–Ω–µ–º—É —Ä–µ–≥–∏—Å—Ç—Ä—É
    
    # –†–µ–≥—É–ª—è—Ä–Ω–æ–µ –≤—ã—Ä–∞–∂–µ–Ω–∏–µ –¥–ª—è –ø–æ–∏—Å–∫–∞ –∑–Ω–∞–∫–æ–≤ –ø—Ä–µ–ø–∏–Ω–∞–Ω–∏—è –∏ –ø–æ—Å–ª–µ–¥—É—é—â–µ–≥–æ —Å–ª–æ–≤–∞
    text = re.sub(r'([,:]) (\w+)', process_after_punctuation, text)
        
    if not re.search(r'[.!?]$', text):
        text += '.'

    return text


# —Ä–∞–∑–¥–µ–ª—è–µ—Ç —Ç–µ–∫—Å –ø–æ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è–º, —Ñ—É–Ω–∫—Ü–∏—è –∏—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –≤ working_with_reports.py
def split_sentences_if_needed(text, language=None):
    """
    Splits a sentence into multiple sentences using SpaCy tokenizer.
    Args:
        text (str): The input sentence text.
    Returns:
        tuple: (list of valid sentences, list of excluded sentences).
    """
    logger.info(f"(—Ñ—É–Ω–∫—Ü–∏—è split_sentences_if_needed) -----------------------------------------")
    logger.info(f"(—Ñ—É–Ω–∫—Ü–∏—è split_sentences_if_needed) üöÄ –ù–∞—á–∞—Ç–æ —Ä–∞–∑–±–∏–µ–Ω–∏–µ —Ç–µ–∫—Å—Ç–∞ –Ω–∞ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è")
    logger.info(f"(—Ñ—É–Ω–∫—Ü–∏—è split_sentences_if_needed) –ø–æ–ª—É—á–µ–Ω —Ç–µ–∫—Å—Ç: {text}")
    # –ó–∞–≥—Ä—É–∂–∞–µ–º –º–æ–¥–µ–ª—å SpaCy –¥–ª—è —Ç–µ–∫—É—â–µ–≥–æ —è–∑—ã–∫–∞
    try:
        nlp = SpacyModel.get_instance(language)
    except ValueError as e:
        logger.error(f"(—Ñ—É–Ω–∫—Ü–∏—è split_sentences_if_needed) ‚ùå –û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –º–æ–¥–µ–ª–∏ SpaCy: {e}")
        return [], []
    
    doc = nlp(text)
    sentences = [sent.text for sent in doc.sents]
    splited_sentences = []
    for sentence in sentences:
        if re.search(r'[a-zA-Z–∞-—è–ê-–Ø—ë–Å0-9]', sentence) and len(sentence.strip()) > 2:
            logger.info(f"(—Ñ—É–Ω–∫—Ü–∏—è split_sentences_if_needed) ‚úÖ –¥–æ–±–∞–≤–ª—è—é –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ –≤ —Å–ø–∏—Å–æ–∫: ({sentence})")
            splited_sentences.append(sentence.strip())
        else:
            logger.info(f"(—Ñ—É–Ω–∫—Ü–∏—è split_sentences_if_needed) ‚ùå –ü—Ä–æ–ø—É—Å–∫–∞—é –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ: ({sentence}) - –Ω–µ —Å–æ–¥–µ—Ä–∂–∏—Ç –±—É–∫–≤ –∏–ª–∏ —Å–ª–∏—à–∫–æ–º –∫–æ—Ä–æ—Ç–∫–æ–µ")
    logger.info(f"(—Ñ—É–Ω–∫—Ü–∏—è split_sentences_if_needed) –ü–æ—Å–ª–µ —Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π –º—ã –ø–æ–ª—É—á–∏–ª–∏ -  ({splited_sentences}) ")
    if len(splited_sentences) > 1:
        # –ï—Å–ª–∏ –Ω–∞–π–¥–µ–Ω–æ –±–æ–ª–µ–µ –æ–¥–Ω–æ–≥–æ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è, –¥–æ–±–∞–≤–ª—è–µ–º –≤ –∏—Å–∫–ª—é—á–µ–Ω–∏—è
        return [], splited_sentences # Splitted sentences
    return splited_sentences, [] # Unsplitted sentences



def sort_key_words_group(unsorted_key_words_group):
    """
    –°–æ—Ä—Ç–∏—Ä–æ–≤–∫–∞ –≥—Ä—É–ø–ø –∫–ª—é—á–µ–≤—ã—Ö —Å–ª–æ–≤ –ø–æ –ø–µ—Ä–≤–æ–π –±—É–∫–≤–µ –ø–µ—Ä–≤–æ–≥–æ –∫–ª—é—á–µ–≤–æ–≥–æ —Å–ª–æ–≤–∞ –≤ –≥—Ä—É–ø–ø–µ.
    –†–∞–±–æ—Ç–∞–µ—Ç –¥–ª—è –≤—Å–µ—Ö –≤–∞—Ä–∏–∞–Ω—Ç–æ–≤ —Å—Ç—Ä—É–∫—Ç—É—Ä—ã: —Å –æ—Ç—á–µ—Ç–∞–º–∏, —Å –∏–Ω–¥–µ–∫—Å–∞–º–∏ –∏ –±–µ–∑ –Ω–∏—Ö.
    """
    key_words_group_with_first_letter = []

    for group_data in unsorted_key_words_group:
        # –í –∑–∞–≤–∏—Å–∏–º–æ—Å—Ç–∏ –æ—Ç —Å—Ç—Ä—É–∫—Ç—É—Ä—ã –¥–∞–Ω–Ω—ã—Ö –≤—ã–±–∏—Ä–∞–µ–º –∫–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞
        key_words = group_data.get("key_words", [])

        if key_words:
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –∫–∞–∫ –ø—Ä–µ–¥—Å—Ç–∞–≤–ª–µ–Ω—ã –∫–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞ ‚Äî –∫–∞–∫ —Å—Ç—Ä–æ–∫–∏ –∏–ª–∏ –æ–±—ä–µ–∫—Ç—ã —Å –ø–æ–ª–µ–º "word"
            if isinstance(key_words[0], dict):
                # –ï—Å–ª–∏ –∫–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞ ‚Äî —ç—Ç–æ —Å–ª–æ–≤–∞—Ä–∏, –±–µ—Ä–µ–º –ø–æ–ª–µ "word"
                first_letter = key_words[0]["word"][0].lower() if key_words[0]["word"] else ""
            else:
                # –ï—Å–ª–∏ –∫–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞ ‚Äî —ç—Ç–æ —Å—Ç—Ä–æ–∫–∏
                first_letter = key_words[0][0].lower() if key_words[0] else ""
        else:
            first_letter = ""  # –ï—Å–ª–∏ –∫–ª—é—á–µ–≤—ã—Ö —Å–ª–æ–≤ –Ω–µ—Ç, –∏—Å–ø–æ–ª—å–∑—É–µ–º –ø—É—Å—Ç—É—é —Å—Ç—Ä–æ–∫—É

        # –î–æ–±–∞–≤–ª—è–µ–º –ø–µ—Ä–≤—É—é –±—É–∫–≤—É –∏ —Å–∞–º—É –≥—Ä—É–ø–ø—É –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —Å–æ—Ä—Ç–∏—Ä–æ–≤–∫–∏
        key_words_group_with_first_letter.append((first_letter, group_data))

    # –°–æ—Ä—Ç–∏—Ä—É–µ–º —Å–ø–∏—Å–æ–∫ –ø–æ –ø–µ—Ä–≤–æ–π –±—É–∫–≤–µ
    sorted_key_words_group = sorted(key_words_group_with_first_letter, key=lambda x: x[0])

    # –ò–∑–≤–ª–µ–∫–∞–µ–º —Ç–æ–ª—å–∫–æ –æ—Ç—Å–æ—Ä—Ç–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ –≥—Ä—É–ø–ø—ã (–±–µ–∑ –ø–µ—Ä–≤–æ–π –±—É–∫–≤—ã)
    return [group_data for _, group_data in sorted_key_words_group]


def group_keywords(keywords, with_index=False, with_report=False):
    """
    –°–æ–∑–¥–∞–µ–º —Å–ø–∏—Å–æ–∫ —Å–≥—Ä—É–ø–ø–∏—Ä–æ–≤–∞–Ω–Ω—ã—Ö —Å–ª–æ–≤ –∏–ª–∏ —Å–ø–∏—Å–æ–∫ —Å–ª–æ–≤–∞—Ä–µ–π 
    —Å –∫–ª—é—á–∞–º–∏ group_index –∏ key_words –≤ –∑–∞–≤–∏—Å–∏–º–æ—Å—Ç–∏ –æ—Ç –ø–æ–ª—É—á–µ–Ω–Ω—ã—Ö 
    –Ω–∞ –≤—Ö–æ–¥–µ –∞—Ä–≥—É–º–µ–Ω—Ç–æ–≤, –∞ —Ç–∞–∫–∂–µ —Å–ø–∏—Å–æ–∫ –∫–ª—é—á–µ–≤—ã—Ö —Å–ª–æ–≤ —Å –æ—Ç—á–µ—Ç–∞–º–∏,
    –µ—Å–ª–∏ with_report=True.

    Args:
        keywords (list): –°–ø–∏—Å–æ–∫ –æ–±—ä–µ–∫—Ç–æ–≤ –∫–ª–∞—Å—Å–∞ KeyWordGroup.
        with_index (boolean): –§–ª–∞–≥ –¥–ª—è –æ–ø—Ä–µ–¥–µ–ª–µ–Ω–∏—è —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞ —Å –∏–ª–∏ –±–µ–∑ –¥–æ–±–∞–≤–ª–µ–Ω–∏—è –∏–Ω–¥–µ–∫—Å–∞.
        with_report (boolean): –§–ª–∞–≥ –¥–ª—è –≥—Ä—É–ø–ø–∏—Ä–æ–≤–∫–∏ –∫–ª—é—á–µ–≤—ã—Ö —Å–ª–æ–≤ –ø–æ –æ—Ç—á–µ—Ç–∞–º.
    
    Returns:
    with_index=True:
        list of dict: –°–ø–∏—Å–æ–∫ —Å–ª–æ–≤–∞—Ä–µ–π —Å –∫–ª—é—á–∞–º–∏ group_index –∏ key_words (–≤–∫–ª—é—á–∞–µ—Ç id –∏ —Å–ª–æ–≤–æ).
    with_index=False:
        list of lists: –°–ø–∏—Å–æ–∫ —Å–ø–∏—Å–∫–æ–≤ —Å –æ—Ç—Å–æ—Ä—Ç–∏—Ä–æ–≤–∞–Ω–Ω—ã–º–∏ –ø–æ group_index –∫–ª—é—á–µ–≤—ã–º–∏ —Å–ª–æ–≤–∞–º–∏ (–≤–∫–ª—é—á–∞–µ—Ç id –∏ —Å–ª–æ–≤–æ).
    with_report=True:
        list of dict: –°–ø–∏—Å–æ–∫ —Å–ª–æ–≤–∞—Ä–µ–π, –≥–¥–µ –∫–ª—é—á–∏ report_id, report_name, group_index –∏ key_words (–≤–∫–ª—é—á–∞–µ—Ç id –∏ —Å–ª–æ–≤–æ).
    """
    
    # –ì—Ä—É–ø–ø–∏—Ä—É–µ–º –∫–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞ –ø–æ group_index, –¥–æ–±–∞–≤–ª—è—è –∏—Ö id –∏ —Å–ª–æ–≤–æ
    grouped_keywords = defaultdict(list)
    for keyword in keywords:
        grouped_keywords[keyword.group_index].append({'id': keyword.id, 'word': keyword.key_word})

    if with_report:
        # –ò—Å–ø–æ–ª—å–∑—É–µ–º —Å–ø–∏—Å–æ–∫ –¥–ª—è —Ö—Ä–∞–Ω–µ–Ω–∏—è –¥–∞–Ω–Ω—ã—Ö –ø–æ –æ—Ç—á–µ—Ç–∞–º
        report_key_words = []
        seen_groups = set()  # –ú–Ω–æ–∂–µ—Å—Ç–≤–æ –¥–ª—è —Ö—Ä–∞–Ω–µ–Ω–∏—è —É–Ω–∏–∫–∞–ª—å–Ω—ã—Ö –≥—Ä—É–ø–ø

        for keyword in keywords:
            group_index = keyword.group_index
            group_words = grouped_keywords[group_index]

            # –î–ª—è –∫–∞–∂–¥–æ–≥–æ –æ—Ç—á–µ—Ç–∞, —Å–≤—è–∑–∞–Ω–Ω–æ–≥–æ —Å –∫–ª—é—á–µ–≤—ã–º —Å–ª–æ–≤–æ–º
            for report in keyword.key_word_reports:
                # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –Ω–µ –±—ã–ª–∞ –ª–∏ –≥—Ä—É–ø–ø–∞ —É–∂–µ –¥–æ–±–∞–≤–ª–µ–Ω–∞ –¥–ª—è —ç—Ç–æ–≥–æ –æ—Ç—á–µ—Ç–∞
                if (report.id, group_index) not in seen_groups:
                    # –î–æ–±–∞–≤–ª—è–µ–º –≤—Å—é –≥—Ä—É–ø–ø—É –∫–ª—é—á–µ–≤—ã—Ö —Å–ª–æ–≤, –µ—Å–ª–∏ –µ—ë –µ—â–µ –Ω–µ –¥–æ–±–∞–≤–∏–ª–∏
                    report_key_words.append({
                        'report_id': report.id,
                        'report_name': report.report_name,
                        'group_index': group_index,
                        'key_words': group_words  # –î–æ–±–∞–≤–ª—è–µ–º –≤—Å—é –≥—Ä—É–ø–ø—É –∫–ª—é—á–µ–≤—ã—Ö —Å–ª–æ–≤ —Å –∏—Ö id
                    })
                    # –û—Ç–º–µ—á–∞–µ–º, —á—Ç–æ –≥—Ä—É–ø–ø–∞ —Å —ç—Ç–∏–º group_index —É–∂–µ –¥–æ–±–∞–≤–ª–µ–Ω–∞ –¥–ª—è –¥–∞–Ω–Ω–æ–≥–æ –æ—Ç—á–µ—Ç–∞
                    seen_groups.add((report.id, group_index))

        return report_key_words

    # –ï—Å–ª–∏ with_index=True, –≤–æ–∑–≤—Ä–∞—â–∞–µ–º —Å–ø–∏—Å–æ–∫ —Å–ª–æ–≤–∞—Ä–µ–π —Å group_index
    if with_index:
        grouped_keywords_with_index = []
        for group_index, words in grouped_keywords.items():
            grouped_keywords_with_index.append({'group_index': group_index, 'key_words': words})
        return grouped_keywords_with_index
    
    # –ï—Å–ª–∏ with_index=False, –≤–æ–∑–≤—Ä–∞—â–∞–µ–º –ø—Ä–æ—Å—Ç–æ —Å–≥—Ä—É–ø–ø–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ –∫–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞ —Å –∏—Ö id
    return list(grouped_keywords.values())


# –°—Ä–∞–≤–Ω–∏–≤–∞—é 2 –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è. –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –≤ working_with_report/save_modified_sentences. 
# –ò—â–µ—Ç —Å–æ–≤–ø–∞–¥–µ–Ω–∏—è —Å –∑–∞–¥–∞–Ω–Ω—ã–º –ø–æ—Ä–æ–≥–æ–º, —Ç–∞–∫–∂–µ –æ—á–∏—â–∞–µ—Ç —Ç–µ–∫—Å—Ç –æ—Ç —á–∏—Å–µ–ª –∏ –∫–ª—é—á–µ–≤—ã—Ö —Å–ª–æ–≤
def compare_sentences_by_paragraph(new_sentences, report_id, profile_id=None):    
    """
    Compares new sentences with existing sentences in their respective paragraphs to determine uniqueness.
    """
    logger.info(f"(—Ñ—É–Ω–∫—Ü–∏—è compare_sentences_by_paragraph) üöÄ –ù–∞—á–∞—Ç–æ —Å—Ä–∞–≤–Ω–µ–Ω–∏–µ –Ω–æ–≤—ã—Ö –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π —Å —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–º–∏ –≤ –±–∞–∑–µ –¥–∞–Ω–Ω—ã—Ö")
    logger.debug(f"(—Ñ—É–Ω–∫—Ü–∏—è compare_sentences_by_paragraph) –ü–æ–ª—É—á–µ–Ω—ã –Ω–æ–≤—ã–µ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è - ({new_sentences})")
    similarity_threshold_fuzz = int(AppConfig.get_setting(profile_id, "SIMILARITY_THRESHOLD_FUZZ", 80))
    except_words = AppConfig.get_setting(profile_id, "EXCEPT_WORDS", "").split(",")
    logger.debug(f"(—Ñ—É–Ω–∫—Ü–∏—è compare_sentences_by_paragraph) –ü–æ—Ä–æ–≥ —Å—Ö–æ–∂–µ—Å—Ç–∏: {similarity_threshold_fuzz}")
    logger.info(f"(—Ñ—É–Ω–∫—Ü–∏—è compare_sentences_by_paragraph) –ò—Å–∫–ª—é—á–∞–µ–º—ã–µ —Å–ª–æ–≤–∞: {except_words}")
    
    existing_paragraphs = Paragraph.query.filter_by(report_id=report_id).all()
    key_words_obj = KeyWord.get_keywords_for_report(profile_id, report_id)
    key_words = [keyword.key_word for keyword in key_words_obj]
    logger.debug(f"(—Ñ—É–Ω–∫—Ü–∏—è compare_sentences_by_paragraph) –ü–æ–ª—É—á–µ–Ω–æ {len(key_words)} –∫–ª—é—á–µ–≤—ã—Ö —Å–ª–æ–≤.")
    
    duplicates = []
    unique_sentences = []
    errors_count = 0
    # –ò—Ç–µ—Ä–∏—Ä—É–µ–º –ø–æ –Ω–æ–≤—ã–º –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è–º
    for new_sentence in new_sentences:
        new_paragraph_id = int(new_sentence.get("paragraph_id"))
        new_text = new_sentence.get("text")
        new_sentence_type = new_sentence.get("sentence_type")
        new_sentence_head_sentence_id = new_sentence.get("head_sentence_id") or None
        
        if not new_paragraph_id or not new_text.strip():
            errors_count += 1
            logger.warning(f"(—Ñ—É–Ω–∫—Ü–∏—è compare_sentences_by_paragraph) ‚ö†Ô∏è –ü—Ä–æ–ø—É—Å–∫–∞—é –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ ({new_sentence}) —Å –ø—É—Å—Ç—ã–º —Ç–µ–∫—Å—Ç–æ–º –∏–ª–∏ –ø—É—Å—Ç—ã–º id –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞")
            continue  # –ü—Ä–æ–ø—É—Å–∫–∞–µ–º –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ
        
        # –ù–∞—Ö–æ–¥–∏–º —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤—É—é—â–∏–π –ø–∞—Ä–∞–≥—Ä–∞—Ñ
        paragraph = next((p for p in existing_paragraphs if p.id == new_paragraph_id), None)
        # –ù–∞—Ö–æ–¥–∏–º —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤—É—é—â—É—é –≥—Ä—É–ø–ø—É –¥–ª—è –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è
        related_group_id = None
        sentence_group_class = None
        
        if new_sentence_type == "body":
            head_sentence = HeadSentence.query.get(new_sentence_head_sentence_id)
            if not head_sentence:
                logger.warning(f"(—Ñ—É–Ω–∫—Ü–∏—è compare_sentences_by_paragraph) ‚ö†Ô∏è –ì–ª–∞–≤–Ω–æ–µ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ —Å id={new_sentence_head_sentence_id} –Ω–µ –Ω–∞–π–¥–µ–Ω–æ. –ü—Ä–æ–ø—É—Å–∫–∞—é –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ ({new_sentence})")
                errors_count += 1
                continue
            related_group_id = head_sentence.body_sentence_group_id or None
            existing_sentences = BodySentenceGroup.get_group_sentences(related_group_id) or []
            logger.debug(f"(—Ñ—É–Ω–∫—Ü–∏—è compare_sentences_by_paragraph) {existing_sentences}")
            existing_sentences.append({"id": head_sentence.id, "sentence": head_sentence.sentence})
            
        else:
            if not paragraph:
                logger.warning(f"(—Ñ—É–Ω–∫—Ü–∏—è compare_sentences_by_paragraph) ‚ö†Ô∏è –ü–∞—Ä–∞–≥—Ä–∞—Ñ —Å id={new_paragraph_id} –Ω–µ –Ω–∞–π–¥–µ–Ω. –ü—Ä–æ–ø—É—Å–∫–∞—é –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ ({new_sentence})")
                errors_count += 1
                continue
            related_group_id = paragraph.tail_sentence_group_id or None
            if not related_group_id:
                logger.warning(f"(—Ñ—É–Ω–∫—Ü–∏—è compare_sentences_by_paragraph) –ì—Ä—É–ø–ø–∞ —Ö–≤–æ—Å—Ç–æ–≤—ã—Ö –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π –Ω–µ –Ω–∞–π–¥–µ–Ω–∞. –î–æ–±–∞–≤–ª—è—é –≤ —É–Ω–∏–∫–∞–ª—å–Ω—ã–µ")
                unique_sentences.append(new_sentence)
                continue
            existing_sentences = TailSentenceGroup.get_group_sentences(related_group_id)
            
            logger.debug(f"(—Ñ—É–Ω–∫—Ü–∏—è compare_sentences_by_paragraph) {existing_sentences}")
        
        # –û—á–∏—â–∞–µ–º —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–µ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è
        cleaned_existing = []
        for sent in existing_sentences:
            cleaned_item = {
                "id": sent.get("id"),
                "original_text": sent["sentence"],
                "cleaned_text": clean_text_with_keywords(sent.get("sentence"), key_words, except_words)
            }
            cleaned_existing.append(cleaned_item)
            
        # –û—á–∏—â–∞–µ–º –Ω–æ–≤–æ–µ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ
        cleaned_new_text = clean_text_with_keywords(new_text, key_words, except_words)
        # –ü—Ä–æ–≤–µ—Ä—è–µ–º –Ω–∞ —Å—Ö–æ–∂–µ—Å—Ç—å —Å —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–º–∏ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è–º–∏
        is_duplicate = False
        for existing in cleaned_existing:
            similarity_rapidfuzz = fuzz.ratio(cleaned_new_text, existing["cleaned_text"])
            if similarity_rapidfuzz >= similarity_threshold_fuzz:
                duplicates.append({
                    "new_sentence": new_sentence,
                   
                    "new_sentence_paragraph": new_sentence.get("paragraph_id"),
                    "matched_with": {
                                "id": existing["id"],
                                "text": existing["original_text"]
                            },
                    "similarity_rapidfuzz": similarity_rapidfuzz
                })
                is_duplicate = True
                logger.debug(f"(—Ñ—É–Ω–∫—Ü–∏—è compare_sentences_by_paragraph) üîÑ –ù–∞–π–¥–µ–Ω–æ –¥—É–±–ª–∏—Ä—É—é—â–µ–µ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ ({new_sentence}) —Å —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–º ({existing['original_text']}) —Å –ø–æ—Ö–æ–∂–µ—Å—Ç—å—é {similarity_rapidfuzz}")
                break

        if not is_duplicate:
            unique_sentences.append(new_sentence)

    return {"duplicates": duplicates, "unique": unique_sentences, "errors_count": errors_count}



# –§—É–Ω–∫—Ü–∏—è –¥–ª—è –ø–æ–∏—Å–∫–∞ —Å—É—â–µ—Å—Ç–≤—É—é—â–∏—Ö –∞–Ω–∞–ª–æ–≥–∏—á–Ω—ã—Ö –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π —Ç–æ–≥–æ –∂–µ —Ç–∏–ø–∞ –≤ –±–∞–∑–µ –¥–∞–Ω–Ω—ã—Ö 
# –∏—Å–ø–æ–ª—å–∑—É—é –≤ models.py. –ò—â–µ—Ç 100% —Å–æ–≤–ø–∞–¥–µ–Ω–∏—è
def find_similar_exist_sentence(sentence_text, sentence_type, report_global_modality_id):
    """
    –ò—â–µ—Ç –ø–æ—Ö–æ–∂–∏–µ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è –≤ –±–∞–∑–µ –¥–∞–Ω–Ω—ã—Ö.
    """
    user_id = current_user.id
    tags = None
    logger.info(f"(—Ñ—É–Ω–∫—Ü–∏—è find_similar_exist_sentence)(—Ç–∏–ø –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è: '{sentence_type}') üöÄ –ù–∞—á–∞—Ç –ø–æ–∏—Å–∫ —Å—É—â–µ—Å—Ç–≤—É—é—â–∏—Ö –∞–Ω–∞–ª–æ–≥–∏—á–Ω—ã—Ö –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π –≤ –±–∞–∑–µ –¥–∞–Ω–Ω—ã—Ö")
    
    # –ü–æ–ª—É—á–∞–µ–º –≤—Å–µ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è —Ç–æ–≥–æ –∂–µ —Ç–∏–ø–∞ –∏ —Å —Ç–∞–∫–∏–º–∏ –∂–µ –±–∞–∑–æ–≤—ã–º–∏ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞–º–∏ –∏–∑ –±–∞–∑—ã –¥–∞–Ω–Ω—ã—Ö
    if sentence_type == "head":
        similar_type_sentences = HeadSentence.query.filter_by(tags=tags, report_global_modality_id=report_global_modality_id, user_id = user_id).all()
    elif sentence_type == "body":
        similar_type_sentences = BodySentence.query.filter_by(tags=tags, report_global_modality_id=report_global_modality_id, user_id = user_id).all()
    elif sentence_type == "tail":
        similar_type_sentences = TailSentence.query.filter_by(tags=tags, report_global_modality_id=report_global_modality_id, user_id = user_id).all()
    else:
        raise ValueError(f"Invalid sentence type: {sentence_type}")

    logger.info(f"-----—Ç–∏–ø {sentence_type}------–Ω–∞–π–¥–µ–Ω–æ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π –¥–ª—è –¥–∞–Ω–Ω–æ–≥–æ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è {len(similar_type_sentences)} --------------------")
    # –°—Ä–∞–≤–Ω–∏–≤–∞–µ–º –≤—Ö–æ–¥–Ω–æ–µ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ —Å –∫–∞–∂–¥—ã–º –∏–∑ —Å—É—â–µ—Å—Ç–≤—É—é—â–∏—Ö
    for exist_sentence in similar_type_sentences:
        if exist_sentence.sentence == sentence_text:
            logger.info(f"(—Ñ—É–Ω–∫—Ü–∏—è find_similar_exist_sentence) üß© –ù–∞–π–¥–µ–Ω–æ —Å–æ–≤–ø–∞–¥–µ–Ω–∏–µ —Å –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ–º '{exist_sentence.sentence}' –≤ –±–∞–∑–µ –¥–∞–Ω–Ω—ã—Ö. –í–æ–∑–≤—Ä–∞—â–∞—é –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ")
            return exist_sentence
    logger.info(f"(—Ñ—É–Ω–∫—Ü–∏—è find_similar_exist_sentence) –°–æ–≤–ø–∞–¥–µ–Ω–∏–π –Ω–µ –Ω–∞–π–¥–µ–Ω–æ. –í–æ–∑–≤—Ä–∞—â–∞—é None")
    return None
    
      
      
# –§—É–Ω–∫—Ü–∏—è –¥–ª—è –æ—Ü–µ–Ω–∫–∏ —É–Ω–∏–∫–∞–ª—å–Ω–æ—Å—Ç–∏ –∏–Ω–¥–µ–∫—Å–æ–≤ –≤ –≥–ª–∞–≤–Ω—ã—Ö –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è—Ö –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞. –ù–æ–º–µ—Ä–∞ –∏–Ω–¥–µ–∫—Å–æ–≤ –Ω–µ –¥–æ–ª–∂–Ω—ã –ø–æ–≤—Ç–æ—Ä—è—Ç—å—Å—è
def check_head_sentence_indexes(paragraph_id):
    """
    –ü—Ä–æ–≤–µ—Ä—è–µ—Ç —É–Ω–∏–∫–∞–ª—å–Ω–æ—Å—Ç—å –∏–Ω–¥–µ–∫—Å–æ–≤ –≥–ª–∞–≤–Ω—ã—Ö –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π –≤ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–µ.
    """
    head_sentences_groupe, _ = Paragraph.get_paragraph_groups(paragraph_id)
    
    head_sentences = HeadSentenceGroup.get_sentences(head_sentences_groupe)
    indexes = [sent.index for sent in head_sentences]
    duplicates = [index for index in indexes if indexes.count(index) > 1]
    return duplicates

        

# –§—É–Ω–∫—Ü–∏—è –ø—Ä–æ–≤–µ—Ä—è–µ—Ç –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ –Ω–∞ —É–Ω–∏–∫–∞–ª—å–Ω–æ—Å—Ç—å –∏ –¥–æ–±–∞–≤–ª—è–µ—Ç –µ–≥–æ 
# –≤ —Ä–µ–∑—É–ª—å—Ç–∞—Ç, –µ—Å–ª–∏ –æ–Ω–æ —É–Ω–∏–∫–∞–ª—å–Ω–æ. –†–∞–±–æ—Ç–∞–µ—Ç —Å–æ —Å–ø–∏—Å–∫–æ–º –ø–æ —Å—Å—ã–ª–∫–µ, 
# –Ω–∏—á–µ–≥–æ –Ω–µ –≤–æ–∑–≤—Ä–∞—â–∞–µ—Ç, –ø—Ä–æ—Å—Ç–æ –º–µ–Ω—è–µ—Ç –ø—Ä–µ–¥–æ—Å—Ç–∞–≤–ª–µ–Ω–Ω—ã–µ —Å–ø–∏—Å–æ–∫ –∏ –º–Ω–æ–∂–µ—Å—Ç–≤–æ
def _add_if_unique(raw_text, key_words, except_words, cleaned_list, result_set, threshold):
    """
    –ü—Ä–æ–≤–µ—Ä—è–µ—Ç, —è–≤–ª—è–µ—Ç—Å—è –ª–∏ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ —É–Ω–∏–∫–∞–ª—å–Ω—ã–º, –∏ –¥–æ–±–∞–≤–ª—è–µ—Ç –µ–≥–æ –≤ —Ä–µ–∑—É–ª—å—Ç–∞—Ç.

    Args:
        raw_text (str): –û—Ä–∏–≥–∏–Ω–∞–ª—å–Ω—ã–π —Ç–µ–∫—Å—Ç.
        key_words (list): –ö–ª—é—á–µ–≤—ã–µ —Å–ª–æ–≤–∞ –¥–ª—è –æ—á–∏—Å—Ç–∫–∏.
        except_words (list): –°–ª–æ–≤–∞-–∏—Å–∫–ª—é—á–µ–Ω–∏—è.
        cleaned_list (list): –£–∂–µ –æ—á–∏—â–µ–Ω–Ω—ã–µ —Ç–µ–∫—Å—Ç—ã.
        result_set (set): –£–Ω–∏–∫–∞–ª—å–Ω—ã–µ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è (—Ä–µ–∑—É–ª—å—Ç–∞—Ç).
        threshold (int): –ü–æ—Ä–æ–≥ —Å—Ö–æ–∂–µ—Å—Ç–∏.
    """
    cleaned = clean_text_with_keywords(raw_text, key_words, except_words)
    for existing in cleaned_list:
        if fuzz.ratio(cleaned, existing) >= threshold:
            return
    cleaned_list.append(cleaned)
    result_set.add(raw_text)
   


def convert_template_json_to_text(template_json: list) -> str:
    """
    –ü—Ä–µ–æ–±—Ä–∞–∑—É–µ—Ç —à–∞–±–ª–æ–Ω –æ—Ç—á–µ—Ç–∞ –∏–∑ JSON-—Å—Ç—Ä—É–∫—Ç—É—Ä—ã (—Å–ø–∏—Å–æ–∫ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ —Å head_sentences)
    –≤ –æ–±—ã—á–Ω—ã–π —Ç–µ–∫—Å—Ç —Å –∑–∞–≥–æ–ª–æ–≤–∫–∞–º–∏ –∏ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è–º–∏ –¥–ª—è –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç–∞.

    Args:
        template_json (list): –°–ø–∏—Å–æ–∫ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ —Å –∫–ª—é—á–∞–º–∏ "paragraph" –∏ "head_sentences".

    Returns:
        str: –ü—Ä–æ—Å—Ç–æ–π —Ç–µ–∫—Å—Ç —Å –∑–∞–≥–æ–ª–æ–≤–∫–∞–º–∏ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ –∏ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è–º–∏ –ø–æ–¥ –Ω–∏–º–∏.
    """
    lines = []
    for block in template_json:
        paragraph = block.get("paragraph", "").strip()
        if paragraph:
            lines.append(f"{paragraph}:")
        for sentence in block.get("head_sentences", []):
            sentence_text = sentence.get("sentence", "").strip()
            if sentence_text:
                lines.append(sentence_text)
        lines.append("")  # –ü—É—Å—Ç–∞—è —Å—Ç—Ä–æ–∫–∞ –º–µ–∂–¥—É –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞–º–∏
    return "\n".join(lines).strip()



def split_report_structure_for_ai(report_data: list) -> tuple:
    """
    –§–æ—Ä–º–∏—Ä—É–µ—Ç –¥–≤–µ —Å—Ç—Ä—É–∫—Ç—É—Ä—ã:
    - skeleton: –ø–æ–ª–Ω–∞—è —Å—Ç—Ä—É–∫—Ç—É—Ä–∞ + id –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ –∏ head_sentences
    - ai_input: —Ç–æ–ª—å–∫–æ editable –ø–∞—Ä–∞–≥—Ä–∞—Ñ—ã (is_active=True, is_additional=False, head_sentences not empty) + "Miscellaneous"
      (–í ai_input –Ω–µ—Ç –∫–ª—é—á–µ–π is_active –∏ is_additional!)

    –ü—Ä–∏–º–µ—Ä:
        skeleton = [
            {
                "id": 1,
                "paragraph": "–û–†–ì–ê–ù–´ –ì–†–£–î–ù–û–ô –ö–õ–ï–¢–ö–ò:",
                "is_active": True,
                "is_additional": False,
                "head_sentences": []
            },
            ...
        ]
        ai_input = [
            {
                "id": 2,
                "paragraph": "–õ–µ–≥–∫–∏–µ",
                "head_sentences": [{"id": 11, "sentence": "–ò–Ω—Ñ–∏–ª—å—Ç—Ä–∞—Ü–∏—è –Ω–µ –≤—ã—è–≤–ª–µ–Ω–∞."}]
            },
            {
                "id": "miscellaneous",
                "paragraph": "Miscellaneous",
                "head_sentences": []
            }
        ]

    Args:
        report_data (list): –°–ø–∏—Å–æ–∫ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ –æ—Ç—á–µ—Ç–∞.
    Returns:
        tuple: (skeleton, ai_input)
    """
    skeleton = []
    ai_input = []

    # Build skeleton
    for para in report_data:
        scel = {
            "id": para["id"],
            "paragraph": para["paragraph"],
            "is_active": para.get("is_active", True),
            "is_additional": para.get("is_additional", False),
            "head_sentences": [
                {
                    "id": hs["id"],
                    "sentence": hs["sentence"]
                }
                for hs in para.get("head_sentences", [])
            ]
        }
        skeleton.append(scel)

    # Build ai_input: only paragraphs with is_active==True, is_additional==False, and head_sentences not empty
    for para in skeleton:
        if (
            para.get("is_active", True)
            and not para.get("is_additional", False)
            and len(para.get("head_sentences", [])) > 0
        ):
            ai_para = dict(para)  # shallow copy is enough, head_sentences is list of dicts
            ai_para.pop("is_active", None)
            ai_para.pop("is_additional", None)
            ai_input.append(ai_para)

    # Add "Miscellaneous" paragraph to ai_input
    misc_paragraph = {
        "id": "miscellaneous",
        "paragraph": "Miscellaneous",
        "head_sentences": []
    }
    ai_input.append(misc_paragraph)

    return skeleton, ai_input


def get_sentences_from_report_for_ai(
    report_data: list[dict[str, any]],
    *,
    include_inactive: bool = False,
    include_additional: bool = False
) -> list[dict[str, str | int]]:
    """
    –ü—Ä–µ–≤—Ä–∞—â–∞–µ—Ç –∏—Å—Ö–æ–¥–Ω—É—é —Å—Ç—Ä—É–∫—Ç—É—Ä—É report_data (—Å–ø–∏—Å–æ–∫ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤)
    –≤ –ø–ª–æ—Å–∫–∏–π —Å–ø–∏—Å–æ–∫ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π-—à–∞–±–ª–æ–Ω–æ–≤ –¥–ª—è –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç–∞, –≤–∏–¥–∞:
        [{"id": <int|str>, "sentence": <str>}, ...]

    - –ë–µ—Ä—ë—Ç —Ç–æ–ª—å–∫–æ –∞–∫—Ç–∏–≤–Ω—ã–µ –ø–∞—Ä–∞–≥—Ä–∞—Ñ—ã (is_active=True)
    - –ò—Å–∫–ª—é—á–∞–µ—Ç –¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–µ (is_additional=True)
    - –ü—Ä–æ–ø—É—Å–∫–∞–µ—Ç –ø—É—Å—Ç—ã–µ –ø–∞—Ä–∞–≥—Ä–∞—Ñ—ã
    - –ù–µ –¥–æ–±–∞–≤–ª—è–µ—Ç 'Miscellaneous'
    """
    logger.info("(get_sentences_from_report_for_ai) üöÄ –ù–∞—á–∞—Ç –ø—Ä–æ—Ü–µ—Å—Å –ø–æ–¥–≥–æ—Ç–æ–≤–∫–∏ —Å–ø–∏—Å–∫–∞ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π –¥–ª—è AI")
    out: list[dict[str, str | int]] = []

    for para in report_data or []:
        is_active = para.get("is_active", True)
        is_additional = para.get("is_additional", False)
        if (not include_inactive and not is_active) or (not include_additional and is_additional):
            continue

        for hs in para.get("head_sentences", []):
            if not isinstance(hs, dict):
                logger.debug(f"(get_sentences_from_report_for_ai) ‚ö†Ô∏è –ü—Ä–æ–ø—É—Å–∫–∞—é –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ: {hs}")
                continue
            if "id" not in hs or "sentence" not in hs:
                logger.debug(f"(get_sentences_from_report_for_ai) ‚ö†Ô∏è –ü—Ä–æ–ø—É—Å–∫–∞—é –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ: {hs}")
                continue
            _id, _sent = hs["id"], hs["sentence"]
            if isinstance(_id, (str, int)) and isinstance(_sent, str):
                out.append({"id": _id, "sentence": _sent})
    print("--------------------------------------------------")
    print(f"–í—ã—Ö–æ–¥ —Å–ø–∏—Å–∫–∞ –¥–ª—è –æ—Ç–ø—Ä–∞–≤–∫–∏ –≤ AI: {out[:100]}")
    print("--------------------------------------------------")
    logger.info(f"(get_sentences_from_report_for_ai) ‚úÖ –ü–æ–¥–≥–æ—Ç–æ–≤–ª–µ–Ω–æ {len(out)} –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π –¥–ª—è AI")
    return out


# –ò—Å–ø–æ–ª—å–∑—É—é –≤ –ª–æ–≥–∏–∫–µ –¥–∏–Ω–∞–º–∏–∫–µ –≤ —Ä–µ–∂–∏–º–µ hard
def replace_head_sentences_with_fuzzy_check(main_data, ai_data, threshold=95):
    """
    –û–±–Ω–æ–≤–ª—è–µ—Ç —Ç–µ–∫—Å—Ç—ã head_sentences –≤ main_data –Ω–∞ –æ—Å–Ω–æ–≤–µ ai_data.
    –°—Ç—Ä—É–∫—Ç—É—Ä–∞ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ –∏ head_sentences –¥–æ–ª–∂–Ω–∞ —Å–æ–≤–ø–∞–¥–∞—Ç—å –ø–æ id.
    –ï—Å–ª–∏ –∫–æ–ª-–≤–æ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ –Ω–µ —Å–æ–≤–ø–∞–¥–∞–µ—Ç ‚Äî –ª–æ–≥–∏—Ä—É–µ–º –∏ –≤–æ–∑–≤—Ä–∞—â–∞–µ–º main_data –∫–∞–∫ –µ—Å—Ç—å.
    –ï—Å–ª–∏ –∑–∞–≥–æ–ª–æ–≤–∫–∏ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ –æ—Ç–ª–∏—á–∞—é—Ç—Å—è –±–æ–ª–µ–µ —á–µ–º –Ω–∞ threshold ‚Äî –±—Ä–æ—Å–∞–µ–º ValueError.
    """
    logger.info("(replace_head_sentences_with_fuzzy_check) üöÄ  –ù–∞—á–∞—Ç –ø—Ä–æ—Ü–µ—Å—Å –∑–∞–º–µ–Ω—ã –≥–ª–∞–≤–Ω—ã—Ö –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π —Å–∏–Ω—Ç–µ–∑–∏—Ä–æ–≤–∞–Ω–Ω—ã–º–∏")
    logger.info(f"(replace_head_sentences_with_fuzzy_check) main_data: {len(main_data)} –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤, ai_data: {len(ai_data)} –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤")

    if not isinstance(main_data, list) or not isinstance(ai_data, list):
        logger.error("(replace_head_sentences_with_fuzzy_check) ‚ùå –í—Ö–æ–¥–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ –Ω–µ —è–≤–ª—è—é—Ç—Å—è —Å–ø–∏—Å–∫–∞–º–∏.")
        raise ValueError("–æ—Å–Ω–æ–≤–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ –∏ –æ–±—Ä–∞–±–æ—Ç–∞–Ω–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ –¥–æ–ª–∂–Ω—ã –±—ã—Ç—å —Å–ø–∏—Å–∫–∞–º–∏ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤.")
    # 1. –ë—ã—Å—Ç—Ä–∞—è —Å–≤–µ—Ä–∫–∞ –∫–æ–ª–∏—á–µ—Å—Ç–≤–∞ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤
    if len(main_data) != len(ai_data):
        logger.error(f"(replace_head_sentences_with_fuzzy_check) ‚ùå –ö–æ–ª–∏—á–µ—Å—Ç–≤–æ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ –Ω–µ —Å–æ–≤–ø–∞–¥–∞–µ—Ç: main={len(main_data)} / ai={len(ai_data)}")
        # –ù–µ —Ä–µ–π–∑–∏–º, –ø—Ä–æ–¥–æ–ª–∂–∞–µ–º –ø–æ–ø—ã—Ç–∫—É –ø–æ–¥—Å—Ç–∞–Ω–æ–≤–∫–∏

    # 2. –ò–Ω–¥–µ–∫—Å–∏—Ä—É–µ–º AI-–ø–∞—Ä–∞–≥—Ä–∞—Ñ—ã –ø–æ id –¥–ª—è –±—ã—Å—Ç—Ä–æ–≥–æ –¥–æ—Å—Ç—É–ø–∞
    ai_paragraphs_by_id = {str(p["id"]): p for p in ai_data}

    # 3. –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º –∫–∞–∂–¥—ã–π –ø–∞—Ä–∞–≥—Ä–∞—Ñ main_data
    for main_par in main_data:
        para_id = str(main_par["id"])
        main_title = main_par.get("paragraph", "").strip()
        ai_par = ai_paragraphs_by_id.get(para_id)
        if not ai_par:
            logger.warning(f"(replace_head_sentences_with_fuzzy_check) –ù–µ –Ω–∞–π–¥–µ–Ω –ø–∞—Ä–∞–≥—Ä–∞—Ñ id={para_id} –≤ AI-–æ—Ç–≤–µ—Ç–µ, –ø—Ä–æ–ø—É—Å–∫–∞—é")
            continue

        ai_title = ai_par.get("paragraph", "").strip()
        ratio = fuzz.ratio(main_title, ai_title)
        if ratio < threshold:
            logger.error(f"(replace_head_sentences_with_fuzzy_check) ‚ùå –ó–∞–≥–æ–ª–æ–≤–æ–∫ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞ '{main_title}' –Ω–µ —Å–æ–≤–ø–∞–¥–∞–µ—Ç —Å AI '{ai_title}' (—Å–æ–≤–ø–∞–¥–µ–Ω–∏–µ {ratio}%)")
            raise ValueError(
                f"–ù–µ—Å–æ–≤–ø–∞–¥–µ–Ω–∏–µ —Ç–µ–∫—Å—Ç–∞ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞ id={para_id}\n"
                f"–û–∂–∏–¥–∞–ª–æ—Å—å: '{main_title}'\n"
                f"–ù–∞–π–¥–µ–Ω–æ:    '{ai_title}'\n"
            )
        # –ò–Ω–¥–µ–∫—Å–∏—Ä—É–µ–º head_sentences –ø–æ id –¥–ª—è –∑–∞–º–µ–Ω—ã
        ai_head_by_id = {str(hs["id"]): hs.get("sentence", "") for hs in ai_par.get("head_sentences", [])}
        for main_hs in main_par.get("head_sentences", []):
            hs_id = str(main_hs["id"])
            if hs_id in ai_head_by_id:
                main_hs["sentence"] = ai_head_by_id[hs_id]
            else:
                logger.warning(f"(replace_head_sentences_with_fuzzy_check) ‚ö†Ô∏è –í AI-–æ—Ç–≤–µ—Ç–µ –Ω–µ –Ω–∞–π–¥–µ–Ω–æ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–µ id={hs_id} –≤ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–µ id={para_id}, —É–¥–∞–ª—è—é –æ—Ä–∏–≥–∏–Ω–∞–ª")
                main_hs["sentence"] = ""
    return main_data


# –ò—Å–ø–æ–ª—å–∑—É—é –≤ –ª–æ–≥–∏–∫–µ –¥–∏–Ω–∞–º–∏–∫–∏ –≤ —Ä–µ–∂–∏–º–µ soft
def build_soft_paragraphs(
    *,
    flat_items: list[dict],
    sorted_parag: list[dict],
    report_id: int,
) -> list[dict]:
    """
    –°–æ–±–∏—Ä–∞–µ—Ç soft-–ø—Ä–µ–¥—Å—Ç–∞–≤–ª–µ–Ω–∏–µ –æ—Ç—á—ë—Ç–∞ –±–µ–∑ –ø–æ—Ö–æ–¥–æ–≤ –≤ –ë–î:
      1) –í –Ω–∞—á–∞–ª–æ ‚Äî ¬´—Ä–æ–¥–Ω—ã–µ¬ª –ø–∞—Ä–∞–≥—Ä–∞—Ñ—ã —Ç–µ–∫—É—â–µ–≥–æ –æ—Ç—á—ë—Ç–∞, —É –∫–æ—Ç–æ—Ä—ã—Ö
         is_active=True, is_additional=True, is_impression=False –∏ –µ—Å—Ç—å head/tail.
      2) –û—Å–Ω–æ–≤–Ω–æ–π —Å–∏–Ω—Ç–µ—Ç–∏—á–µ—Å–∫–∏–π –ø–∞—Ä–∞–≥—Ä–∞—Ñ —Å head-–ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è–º–∏ –∏–∑ flat_items
         –≤ –∏—Å—Ö–æ–¥–Ω–æ–º –ø–æ—Ä—è–¥–∫–µ; –¥–ª—è –∏–∑–≤–µ—Å—Ç–Ω—ã—Ö head –ø–æ–¥—Å—Ç–∞–≤–ª—è—é—Ç—Å—è ¬´—Ä–æ–¥–Ω—ã–µ¬ª body
         –∏–∑ sorted_parag –∫–∞–∫ –µ—Å—Ç—å (–±–µ–∑ —Ç—Ä–∞–Ω—Å—Ñ–æ—Ä–º–∞—Ü–∏–π).
         –í—Å–µ tail –≤—Å–µ—Ö –ù–ï-–∏–º–ø—Ä–µ—Å—Å–∏–æ–Ω–Ω—ã—Ö –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ —Å–∫–ª–∞–¥—ã–≤–∞–µ–º –≤ tail —Å–∏–Ω—Ç–µ—Ç–∏—á–µ—Å–∫–æ–≥–æ.
      3) –ü–∞—Ä–∞–≥—Ä–∞—Ñ ¬´–ó–∞–∫–ª—é—á–µ–Ω–∏–µ¬ª: –±–µ—Ä—ë–º —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–π is_impression=True (–µ—Å–ª–∏ –µ—Å—Ç—å),
         head –∑–∞–º–µ–Ω—è–µ–º –µ–¥–∏–Ω—ã–º –±–ª–æ–∫–æ–º –∏–∑ flat_items —Å is_impression=True,
         tail –∏–∑ –∏—Å—Ö–æ–¥–Ω–æ–≥–æ –∏–º–ø—Ä–µ—Å—Å–∏–æ–Ω–Ω–æ–≥–æ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞ –ø–µ—Ä–µ–Ω–æ—Å–∏–º –≤ —Ä–µ–∑—É–ª—å—Ç–∏—Ä—É—é—â–∏–π.
    –í—Å—è —Å–±–æ—Ä–∫–∞ ‚Äî –≤ –ø–∞–º—è—Ç–∏, –∏—Å—Ö–æ–¥–Ω—ã–µ dict'—ã –∏–∑ sorted_parag –Ω–µ –º—É—Ç–∏—Ä—É–µ–º.
    """

    if not isinstance(flat_items, list) or not isinstance(sorted_parag, list):
        raise ValueError("[soft] –ù–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –≤—Ö–æ–¥–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ: flat_items/sorted_parag")

    # 1) –ë—ã—Å—Ç—Ä—ã–µ –∏–Ω–¥–µ–∫—Å—ã –ø–æ –∏—Å—Ö–æ–¥–Ω—ã–º –¥–∞–Ω–Ω—ã–º
    #    head_proto_by_id: id head -> ¬´—Ä–æ–¥–Ω–æ–π¬ª head-—Å–ª–æ–≤–∞—Ä—å (—Å body –∏ –ø—Ä.)
    #    existing_additional_blocks: ¬´—Ä–æ–¥–Ω—ã–µ¬ª –¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–µ –ø–∞—Ä–∞–≥—Ä–∞—Ñ—ã –≤ –Ω–∞—á–∞–ª–æ –≤—ã–¥–∞—á–∏
    #    impression_source_paragraph: ¬´—Ä–æ–¥–Ω–æ–π¬ª –ø–∞—Ä–∞–≥—Ä–∞—Ñ –∑–∞–∫–ª—é—á–µ–Ω–∏—è (–µ—Å–ª–∏ –µ—Å—Ç—å)
    #    main_tails_pool: –≤—Å–µ tail –∏–∑ –ù–ï-–∏–º–ø—Ä–µ—Å—Å–∏–æ–Ω–Ω—ã—Ö –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ (—É–π–¥—É—Ç –≤ –Ω–æ–≤—ã–π –æ—Å–Ω–æ–≤–Ω–æ–π)
    #    impression_tails_pool: tail –∏–∑ –∏–º–ø—Ä–µ—Å—Å–∏–æ–Ω–Ω–æ–≥–æ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞ (—É–π–¥—É—Ç –≤ –Ω–æ–≤—ã–π –∑–∞–∫–ª—é—á–∏—Ç–µ–ª—å–Ω—ã–π)
    head_proto_by_id: dict[int, dict] = {}
    existing_additional_blocks: list[dict] = []
    impression_source_paragraph: dict | None = None
    main_tails_pool: list[dict] = []
    impression_tails_pool: list[dict] = []
    base_index = len(sorted_parag) + 100 # —á—Ç–æ–±—ã –Ω–∞–≤–µ—Ä–Ω—è–∫–∞ –Ω–µ –ø–µ—Ä–µ—Å–µ–∫–∞—Ç—å—Å—è —Å —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–º–∏
    next_fake_id = -1  # –¥–ª—è —Å–∏–Ω—Ç–µ—Ç–∏—á–µ—Å–∫–∏—Ö id
    
    for p in sorted_parag:
        # –ò–Ω–¥–µ–∫—Å–∏—Ä—É–µ–º head
        for hs in p.get("head_sentences", []):
            hid = hs.get("id")
            if isinstance(hid, int):
                # –°–æ—Ö—Ä–∞–Ω—è–µ–º ¬´—Ä–æ–¥–Ω–æ–π¬ª head-—Å–ª–æ–≤–∞—Ä—å —Ü–µ–ª–∏–∫–æ–º (–Ω–µ –∫–æ–ø–∏—Ä—É–µ–º)
                head_proto_by_id[hid] = hs

        # –°–æ–±–∏—Ä–∞–µ–º —Ö–≤–æ—Å—Ç—ã
        tails = p.get("tail_sentences", []) or []

        if p.get("is_impression", False):
            impression_source_paragraph = p
            if tails:
                impression_tails_pool.extend(tails)
        else:
            # ¬´—Ä–æ–¥–Ω—ã–µ¬ª –¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–µ –ø–∞—Ä–∞–≥—Ä–∞—Ñ—ã, –∫–æ—Ç–æ—Ä—ã–µ –Ω–∞–¥–æ –æ—Ç–¥–∞—Ç—å –ø–µ—Ä–µ–¥ —Å–∏–Ω—Ç–µ—Ç–∏—á–µ—Å–∫–∏–º
            if (
                p.get("is_active", True)
                and p.get("is_additional", False)
                and (p.get("head_sentences") or p.get("tail_sentences"))
            ):
                existing_additional_blocks.append(p)
            # —Ö–≤–æ—Å—Ç—ã –æ–±—ã—á–Ω—ã—Ö –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ ‚Äî –≤ –æ—Å–Ω–æ–≤–Ω–æ–π —Å–∏–Ω—Ç–µ—Ç–∏—á–µ—Å–∫–∏–π
            if tails:
                main_tails_pool.extend(tails)

    # 2) –°–∏–Ω—Ç–µ—Ç–∏—á–µ—Å–∫–∏–π –æ—Å–Ω–æ–≤–Ω–æ–π –ø–∞—Ä–∞–≥—Ä–∞—Ñ
    #    –ò–Ω–¥–µ–∫—Å—ã –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π ‚Äî —á–µ—Ä–µ–∑ enumerate, —Å–∏–Ω—Ç–µ—Ç–∏—á–µ—Å–∫–∏–µ id ‚Äî –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω—ã–µ.
    main_paragraph: dict = {
        "id": next_fake_id,  
        "report_id": report_id,
        "paragraph_index": base_index + 1,          
        "paragraph": "",
        "paragraph_visible": False,
        "title_paragraph": False,
        "bold_paragraph": False,
        "is_impression": False,
        "is_active": True,
        "str_before": False,
        "str_after": False,
        "is_additional": True,
        "comment": "synthesized",
        "paragraph_weight": 1,
        "tags": None,
        "head_sentence_group_id": None,
        "tail_sentence_group_id": None,
        "head_sentences": [],
        "tail_sentences": [],
    }

    next_fake_id -= 1  # –¥–ª—è —Å–∏–Ω—Ç–µ—Ç–∏—á–µ—Å–∫–∏—Ö id
    # –ü—Ä–æ–±–µ–≥–∞–µ–º –≤—Å—ë, —á—Ç–æ –ù–ï –ø–æ–º–µ—á–µ–Ω–æ –∫–∞–∫ impression
    seq = [it for it in flat_items if not bool(it.get("is_impression", False))]

    for idx, it in enumerate(seq, start=1):
        raw_id = it.get("id")
        text = it.get("sentence", "")

        # –í—ã—á–∏—Å–ª—è–µ–º head_id: —á–∏—Å–ª–æ ‚Üí –∫–∞–∫ –µ—Å—Ç—å; —Å—Ç—Ä–æ–∫–∞-—á–∏—Å–ª–æ ‚Üí int; –∏–Ω–∞—á–µ ‚Äî –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω—ã–π —Å–∏–Ω—Ç–µ—Ç–∏—á–µ—Å–∫–∏–π
        if isinstance(raw_id, int):
            head_id = raw_id
        elif isinstance(raw_id, str) and raw_id.isdigit():
            head_id = int(raw_id)
        else:
            head_id = next_fake_id
            next_fake_id -= 1

        # –ë–∞–∑–æ–≤–∞—è –∑–∞–≥–æ—Ç–æ–≤–∫–∞ head-–ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è
        head_payload: dict = {
            "id": head_id,
            "sentence": text,         
            "tags": None,
            "comment": "synthesized",
            "is_linked": False,
            "group_id": None,
            "body_sentences": [],
            "body_sentence_group_id": None,
            "sentence_index": idx,
        }

        # –ï—Å–ª–∏ head –∏–∑–≤–µ—Å—Ç–µ–Ω –ø–æ id, –ø–æ–¥—Ç—è–≥–∏–≤–∞–µ–º ¬´—Ä–æ–¥–Ω—ã–µ¬ª –ø–æ–ª—è –∏–∑ –∏—Å—Ö–æ–¥–Ω–æ–≥–æ head-–ø—Ä–æ—Ç–æ
        if isinstance(head_id, int) and head_id in head_proto_by_id:
            proto = head_proto_by_id[head_id]
            # –ø–µ—Ä–µ–Ω–æ—Å–∏–º group_id –∏ —Å–≤—è–∑—å body-–≥—Ä—É–ø–ø—ã
            head_payload["group_id"] = proto.get("group_id")
            head_payload["body_sentence_group_id"] = proto.get("body_sentence_group_id")
            # –∏ –≤—Å—Ç–∞–≤–ª—è–µ–º ¬´–Ω–∞—Ç–∏–≤–Ω—ã–µ¬ª body –∫–∞–∫ –µ—Å—Ç—å (–±–µ–∑ –∫–æ–ø–∏–π –∏ –ø–µ—Ä–µ–±–æ—Ä–æ–∫)
            head_payload["body_sentences"] = proto.get("body_sentences", []) or []

        main_paragraph["head_sentences"].append(head_payload)

    # –î–æ–±–∞–≤–ª—è–µ–º –≤ –æ—Å–Ω–æ–≤–Ω–æ–π –ø–∞—Ä–∞–≥—Ä–∞—Ñ ¬´—Å–æ–±—Ä–∞–Ω–Ω—ã–µ¬ª —Ö–≤–æ—Å—Ç—ã –∏–∑ –≤—Å–µ—Ö –Ω–µ–∏–º–ø—Ä–µ—Å—Å–∏–æ–Ω–Ω—ã—Ö –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤
    main_paragraph["tail_sentences"] = main_tails_pool

    # 3) –ü–∞—Ä–∞–≥—Ä–∞—Ñ ¬´–ó–∞–∫–ª—é—á–µ–Ω–∏–µ¬ª
    #    –°–æ–±–∏—Ä–∞–µ–º —Ç–µ–∫—Å—Ç –∏–∑ flat_items —Å is_impression=True; –µ—Å–ª–∏ –Ω–µ—Å–∫–æ–ª—å–∫–æ ‚Äî —Å–æ–µ–¥–∏–Ω—è–µ–º —á–µ—Ä–µ–∑ \n
    impression_texts = [it.get("sentence", "") for it in flat_items if bool(it.get("is_impression", False))]
    impression_sentence = ""
    if impression_texts:
        impression_sentence = impression_texts[0] if len(impression_texts) == 1 else "\n".join(impression_texts)
    else:
        impression_sentence = "–ó–∞–∫–ª—é—á–µ–Ω–∏–µ –Ω–µ –±—ã–ª–æ —Å–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω–æ –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏."

    if impression_source_paragraph:
        # –î–µ–ª–∞–µ–º –Ω–æ–≤—ã–π –æ–±—ä–µ–∫—Ç –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞ –∑–∞–∫–ª—é—á–µ–Ω–∏—è (–∏—Å—Ö–æ–¥–Ω—ã–π –Ω–µ —Ç—Ä–æ–≥–∞–µ–º)
        impression_paragraph = {
            "id": impression_source_paragraph.get("id", next_fake_id),
            "report_id": report_id,
            "paragraph_index": base_index + 2,
            "paragraph": impression_source_paragraph.get("paragraph", "–ó–∞–∫–ª—é—á–µ–Ω–∏–µ"),
            "paragraph_visible": impression_source_paragraph.get("paragraph_visible", True),
            "title_paragraph": impression_source_paragraph.get("title_paragraph", True),
            "bold_paragraph": impression_source_paragraph.get("bold_paragraph", True),
            "is_impression": True,
            "is_active": impression_source_paragraph.get("is_active", True),
            "str_before": impression_source_paragraph.get("str_before", True),
            "str_after": impression_source_paragraph.get("str_after", True),
            "is_additional": impression_source_paragraph.get("is_additional", False),
            "comment": impression_source_paragraph.get("comment"),
            "paragraph_weight": impression_source_paragraph.get("paragraph_weight", 1),
            "tags": impression_source_paragraph.get("tags"),
            "head_sentence_group_id": None,
            "tail_sentence_group_id": impression_source_paragraph.get("tail_sentence_group_id"),
            "head_sentences": [],
            "tail_sentences": impression_tails_pool,  # ¬´—Ä–æ–¥–Ω—ã–µ¬ª —Ö–≤–æ—Å—Ç—ã –∑–∞–∫–ª—é—á–µ–Ω–∏—è
        }
        next_fake_id -= 1
    else:
        # –ï—Å–ª–∏ –∏—Å—Ö–æ–¥–Ω–æ–≥–æ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞ –∑–∞–∫–ª—é—á–µ–Ω–∏—è –Ω–µ—Ç ‚Äî —Å–∏–Ω—Ç–µ–∑–∏—Ä—É–µ–º –∑–∞–≥–ª—É—à–∫—É
        impression_paragraph = {
            "id": next_fake_id,  
            "report_id": report_id,
            "paragraph_index": base_index + 2,
            "paragraph": "–ó–∞–∫–ª—é—á–µ–Ω–∏–µ",
            "paragraph_visible": True,
            "title_paragraph": True,
            "bold_paragraph": True,
            "is_impression": True,
            "is_active": True,
            "str_before": True,
            "str_after": True,
            "is_additional": False,
            "comment": "synthesized",
            "paragraph_weight": 1,
            "tags": None,
            "head_sentence_group_id": None,
            "tail_sentence_group_id": None,
            "head_sentences": [],
            "tail_sentences": impression_tails_pool,
        }
        next_fake_id -= 1

    # –í—Å—Ç–∞–≤–ª—è–µ–º –µ–¥–∏–Ω—ã–π head —Å impression-—Ç–µ–∫—Å—Ç–æ–º
    impression_head = {
        "id": next_fake_id,         
        "sentence": impression_sentence,
        "tags": None,
        "comment": "synthesized",
        "is_linked": False,
        "group_id": None,
        "body_sentences": [],
        "body_sentence_group_id": None,
        "sentence_index": 1,
    }
    next_fake_id -= 1
    impression_paragraph["head_sentences"].append(impression_head)

    # 4) –§–æ—Ä–º–∏—Ä—É–µ–º –∏—Ç–æ–≥–æ–≤—ã–π —Å–ø–∏—Å–æ–∫ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤:
    #    —Å–Ω–∞—á–∞–ª–∞ ¬´—Ä–æ–¥–Ω—ã–µ¬ª –¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–µ, –∑–∞—Ç–µ–º –Ω–∞—à –æ—Å–Ω–æ–≤–Ω–æ–π —Å–∏–Ω—Ç–µ—Ç–∏—á–µ—Å–∫–∏–π, –∑–∞—Ç–µ–º –∑–∞–∫–ª—é—á–µ–Ω–∏–µ
    paragraphs_out: list[dict] = []
    if existing_additional_blocks:
        paragraphs_out.extend(existing_additional_blocks)
    paragraphs_out.append(main_paragraph)
    paragraphs_out.append(impression_paragraph)

    return paragraphs_out


def merge_ai_response_into_skeleton(skeleton, ai_response):
    """
    –í–æ—Å—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ—Ç –ø–æ–ª–Ω—É—é —Å—Ç—Ä—É–∫—Ç—É—Ä—É –æ—Ç—á–µ—Ç–∞:
    - –û—Å–Ω–æ–≤–Ω—ã–µ –ø–∞—Ä–∞–≥—Ä–∞—Ñ—ã: id, paragraph, is_active, is_additional, head_sentences —Å –Ω–æ–≤—ã–º–∏ —Ç–µ–∫—Å—Ç–∞–º–∏
    - Miscellaneous –≤–æ–∑–≤—Ä–∞—â–∞–µ—Ç—Å—è –æ—Ç–¥–µ–ª—å–Ω–æ (—Å–ø–∏—Å–æ–∫ head_sentences)
    """
    # –ë—ã—Å—Ç—Ä–∞—è –∏–Ω–¥–µ–∫—Å–∞—Ü–∏—è AI-–ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ –ø–æ id
    ai_paragraphs_by_id = {str(p["id"]): p for p in ai_response if str(p["id"]) != "miscellaneous"}
    misc_sentences = []
    for p in ai_response:
        if str(p["id"]) == "miscellaneous":
            misc_sentences = p.get("head_sentences", [])
            break

    # –ò–¥–µ–º –ø–æ —Å–∫–µ–ª–µ—Ç—É –∏ –ø–æ–¥–º–µ–Ω—è–µ–º –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è
    merged = []
    for para in skeleton:
        para_id = str(para["id"])
        # –ö–æ–ø–∏—Ä—É–µ–º –ø–∞—Ä–∞–≥—Ä–∞—Ñ –∏–∑ —Å–∫–µ–ª–µ—Ç–∞
        merged_para = dict(para)
        # –û–±–Ω–æ–≤–ª—è–µ–º —Ç–æ–ª—å–∫–æ —Ç–µ–∫—Å—Ç –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏–π
        ai_para = ai_paragraphs_by_id.get(para_id)
        if ai_para:
            # –ú–∞–ø–ø–∏–Ω–≥ head_sentences –ø–æ id –¥–ª—è —Å–∫–æ—Ä–æ—Å—Ç–∏
            ai_head_by_id = {str(hs["id"]): hs["sentence"] for hs in ai_para.get("head_sentences", [])}
            # –û–±–Ω–æ–≤–ª—è–µ–º –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è —Ç–æ–ª—å–∫–æ –ø–æ id (–Ω–µ –º–µ–Ω—è—è —Å—Ç—Ä—É–∫—Ç—É—Ä—É)
            new_head_sentences = []
            for hs in merged_para["head_sentences"]:
                hs_id = str(hs["id"])
                if hs_id in ai_head_by_id:
                    # –û–±–Ω–æ–≤–ª—è–µ–º —Ç–æ–ª—å–∫–æ —Ç–µ–∫—Å—Ç
                    new_hs = dict(hs)
                    new_hs["sentence"] = ai_head_by_id[hs_id]
                    new_head_sentences.append(new_hs)
                else:
                    # –ù–µ –Ω–∞–π–¥–µ–Ω–æ –≤ AI ‚Äî –æ—Å—Ç–∞–≤–ª—è–µ–º –æ—Ä–∏–≥–∏–Ω–∞–ª
                    new_head_sentences.append(hs)
            merged_para["head_sentences"] = new_head_sentences
        # –ï—Å–ª–∏ –Ω–µ—Ç ai_para ‚Äî –Ω–∏—á–µ–≥–æ –Ω–µ —Ç—Ä–æ–≥–∞–µ–º, –∫–æ–ø–∏—è —Å–∫–µ–ª–µ—Ç–∞
        merged.append(merged_para)
    return merged, misc_sentences


# –ü—Ä–µ–≤—Ä–∞—â–∞–µ—Ç JSON-—Å—Ç—Ä–æ–∫–∏ –≤ –æ–±—ä–µ–∫—Ç—ã Python, –µ—Å–ª–∏ –æ–Ω–∏ —Å–µ—Ä–∏–∞–ª–∏–∑–æ–≤–∞–Ω—ã –ù–ò–ì–î–ï –ù–ï –ò–°–ü–û–õ–¨–ó–£–Æ–¢–°–Ø
def deep_json_deserialize_if_needed(data, context="root"):
    """
    –†–µ–∫—É—Ä—Å–∏–≤–Ω–æ –¥–µ—Å–µ—Ä–∏–∞–ª–∏–∑—É–µ—Ç —Å—Ç—Ä–æ–∫–∏ –Ω–∞ –≤—Å–µ—Ö —É—Ä–æ–≤–Ω—è—Ö, –µ—Å–ª–∏ –æ–Ω–∏ —è–≤–ª—è—é—Ç—Å—è JSON-—Å—Ç—Ä–æ–∫–æ–π
    (–Ω–∞–ø—Ä–∏–º–µ—Ä, —Å–ø–∏—Å–æ–∫ –∏–ª–∏ —Å–ª–æ–≤–∞—Ä—å, —Å–µ—Ä–∏–∞–ª–∏–∑–æ–≤–∞–Ω–Ω—ã–π –¥–≤–∞–∂–¥—ã).
    """
    # 1. –ü–æ–ø—ã—Ç–∫–∞ —Ä–∞—Å–ø–∞—Ä—Å–∏—Ç—å JSON, –µ—Å–ª–∏ —ç—Ç–æ —Å—Ç—Ä–æ–∫–∞
    logger.info(f"(deep_json_deserialize_if_needed) üöÄ –ù–∞—á–∞—Ç–∞ –¥–µ—Å–µ—Ä–∏–∞–ª–∏–∑–∞—Ü–∏—è –¥–∞–Ω–Ω—ã—Ö –≤ –∫–æ–Ω—Ç–µ–∫—Å—Ç–µ: {context}")
    if isinstance(data, str):
        try:
            parsed = json.loads(data)
            logger.warning(f"(deep_json_deserialize_if_needed) ‚ö†Ô∏è –í '{context}' –¥–µ—Å–µ—Ä–∏–∞–ª–∏–∑–æ–≤–∞–Ω–∞ —Å—Ç—Ä–æ–∫–∞ –≤ JSON.")
            return deep_json_deserialize_if_needed(parsed, context)
        except json.JSONDecodeError:
            logger.info(f"(deep_json_deserialize_if_needed) ‚ùå –î–∞–Ω–Ω—ã–µ –≤ '{context}' –Ω–µ —è–≤–ª—è—é—Ç—Å—è JSON-—Å—Ç—Ä–æ–∫–æ–π, –≤–æ–∑–≤—Ä–∞—â–∞—é –∫–∞–∫ –µ—Å—Ç—å.")
            return data  # –æ–±—ã—á–Ω–∞—è —Å—Ç—Ä–æ–∫–∞, –Ω–µ JSON

    # 2. –û–±—Ä–∞–±–æ—Ç–∫–∞ —Å–ª–æ–≤–∞—Ä—è
    elif isinstance(data, dict):
        return {
            key: deep_json_deserialize_if_needed(value, context + f".{key}")
            for key, value in data.items()
        }

    # 3. –û–±—Ä–∞–±–æ—Ç–∫–∞ —Å–ø–∏—Å–∫–∞
    elif isinstance(data, list):
        return [
            deep_json_deserialize_if_needed(item, context + f"[{i}]")
            for i, item in enumerate(data)
        ]

    # 4. –í—Å—ë –æ—Å—Ç–∞–ª—å–Ω–æ–µ –≤–æ–∑–≤—Ä–∞—â–∞–µ–º –∫–∞–∫ –µ—Å—Ç—å (int, float, None –∏ –ø—Ä.)
    return data

