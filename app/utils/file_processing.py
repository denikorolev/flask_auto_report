# file_processing.py

from flask import session
from flask_security import current_user
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename
from docx import Document
import os
from unidecode import unidecode
from datetime import datetime
import tempfile
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from config import get_config, Config
from app.models.models import db, FileMetadata, Report, HeadSentenceGroup, TailSentenceGroup, BodySentenceGroup, KeyWord, ReportSubtype
from app.utils.sentence_processing import clean_text_with_keywords, _add_if_unique
from openai import OpenAI
from app.utils.logger import logger
import easyocr
from app.utils.redis_client import redis_get, redis_set, redis_delete


# Проверка допустимости расширения загружаемого файла
def allowed_file(file_name, file_type):
    if file_type == "doc":
        return '.' in file_name and file_name.rsplit('.', 1)[1].lower() in {'docx', "doc", "odt"}
    elif file_type == "tab":
        return '.' in file_name and file_name.rsplit('.', 1)[1].lower() in {"xlsx"}
    elif file_type == "jpg":
        return '.' in file_name and file_name.rsplit('.', 1)[1].lower() in {"jpg", "jpeg", "png"}
    elif file_type == "json":
        return '.' in file_name and file_name.rsplit('.', 1)[1].lower() in {"json"}
    else:
        return False

# Преобразование пути к файлу в объект файла
def create_filestorage_from_path(file_path):
    """
    Преобразует файл с диска в объект FileStorage, который Flask использует для загруженных файлов.
    
    :param file_path: Путь к файлу на диске.
    :return: Объект FileStorage.
    """
    try:
        # Открываем файл в режиме чтения
        file = open(file_path, "rb")
        # Создаем объект FileStorage, передавая файл и его имя
        file_storage = FileStorage(file, filename=os.path.basename(file_path))
        return file_storage
    except Exception as e:
        logger.error(f"Error creating FileStorage from path {file_path}: {e}")
        return None

def sanitize_filename(filename):
    """
    Удаляет пробелы, тире и выполняет транслитерацию имени файла.
    
    :param filename: Исходное имя файла.
    :return: Обработанное имя файла.
    """
    # Заменяем пробелы и тире на подчеркивания
    sanitized_name = filename.replace(" ", "_").replace("-", "_")
    # Выполняем транслитерацию (убираем специальные символы)
    return unidecode(sanitized_name)


def sync_profile_files(profile_id, user_id, user_email):
    """
    Синхронизирует файлы профиля: проверяет записи в базе данных и соответствующие файлы на диске.
    Удаляет записи в базе данных, если файла нет на диске, и удаляет файлы, если нет записи в базе данных.
    
    :param profile_id: ID профиля, для которого проводится синхронизация.
    """
    try:
        # Получаем папку профиля
        upload_folder = Config.get_user_upload_folder(user_id, profile_id, user_email)

        # 1. Проверяем наличие всех файлов из базы данных
        files_in_db = FileMetadata.query.filter_by(profile_id=profile_id).all()
        
        for file_meta in files_in_db:
            file_path = file_meta.file_path

            if not os.path.exists(file_path):
                db.session.delete(file_meta)
                db.session.commit()

        # 2. Проверяем наличие всех файлов в папке профиля
        for root, dirs, files in os.walk(upload_folder):
            for file in files:
                file_path = os.path.join(root, file)

                # Проверяем, есть ли этот файл в базе данных
                file_record = FileMetadata.query.filter_by(file_path=file_path).first()

                if not file_record:
                    os.remove(file_path)

        return "Synchronization completed successfully."

    except Exception as e:
        logger.error(f"Error during file synchronization: {e}")
        return f"Error during file synchronization: {e}"


def file_uploader(file, file_type, folder_name, file_name=None, file_description=None, user_id=None, profile_id=None, user_email=None):
    """
    Uploads a file to the server in the folder corresponding to the user's profile and specified folder name,
    and saves metadata to the database.
    
    :param file: файл, который загружается или путь к файлу
    :param file_type: тип файла (например, "doc" или "jpg")
    :param folder_name: имя папки, в которую будет загружен файл (например, "templates" или "word_reports")
    :param file_name: имя файла (если не передано, используется имя папки)
    :param file_description: описание файла для записи в базу данных
    
    :return: сообщение об успехе или ошибке и путь к файлу.
    """
    # Проверяем, является ли file строкой, если да, то это путь к файлу
    if isinstance(file, str):
        file = create_filestorage_from_path(file)
        if not file:
            return "Failed to convert file path to FileStorage", None

    # Проверка, есть ли у файла атрибут filename
    if not hasattr(file, "filename") or file.filename == "":
        return "The file name is empty", None

    # Проверка допустимого расширения файла
    if not allowed_file(file.filename, file_type):
        return "The file name or extension is not allowed", None
    
    # Получаем расширение файла
    name, ext = os.path.splitext(file.filename)
    
    # Если имя файла не передано, используем имя папки в качестве имени файла
    if not file_name:
        file_name = folder_name
    sanitized_name = sanitize_filename(file_name)
    # Формируем строку с текущей датой и временем
    date_str = datetime.now().strftime("%d_%m_%y")
    time_str = datetime.now().strftime("%H_%M")
    
    # Создаем имя файла с датой и временем
    filename_with_date_time = f"{sanitized_name}_{date_str}_{time_str}{ext}"
    
    # Получаем путь к папке пользователя с учетом текущего профиля
    try:
        user_folder = get_config().get_user_upload_folder(user_id, profile_id, user_email)
    except Exception as e:
        logger.error(f"Failed to get user folder: {e}")
        return f"Failed to get user folder: {e}", None
    
    # Путь к папке, соответствующей folder_name и текущей дате
    target_folder = os.path.join(user_folder, folder_name, f"{folder_name}_{date_str}")
    
    # Создаем папку, если она не существует
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
    
    # Полный путь к файлу
    file_path = os.path.join(target_folder, filename_with_date_time)
    
    try:
        # Сохраняем файл по указанному пути
        file.save(file_path)

        # Создаем запись о файле в базе данных
        file_description = file_description or folder_name
        try:
            FileMetadata.create(
                profile_id=profile_id,
                file_name=filename_with_date_time,
                file_path=file_path,
                file_type=file_type,
                file_description=file_description
            )
        except Exception as e:
            logger.error(f"Error saving file metadata: {e}")
            return f"Can't add data to file metadata: {e}", None
        
        return "The file was uploaded successfully and metadata saved", file_path
    
    except Exception as e:
        logger.error(f"Error uploading file: {e}")
        return f"The file wasn't uploaded due to the following error: {e}", None


def save_to_word(text, name, subtype, report_type, birthdate, reportnumber, scanParam, side=""):
    # Убираем лишние пробелы и проверяем на пустые строки
    name = name.strip() or "NoName"
    subtype = subtype.strip() or "NoSubtype"
    report_type = report_type.strip() or "NoReportType"
    profile_id = session.get("profile_id")

    try:
        date_str = datetime.now().strftime("%d_%m_%y")
        modified_date_str = date_str.replace("_", ".")
        # Преобразование строки в объект datetime
        trans_birthdate = datetime.strptime(birthdate, "%Y-%m-%d")
        # Форматирование даты
        modified_birthdate = trans_birthdate.strftime("%d.%m.%y")
    except:
        logger.error("Error parsing birthdate, setting to 'unknown'")
        modified_birthdate = "unknown"
    try:
        profile_id = profile_id
        user_id = current_user.id

        signatura_metadata = FileMetadata.get_file_by_description(profile_id, "signatura")
        template_metadata = FileMetadata.get_file_by_description(profile_id, "word_template")
        
        if not template_metadata:
            raise Exception("Word template not found in the database.")
        
        signatura_path = signatura_metadata.file_path
        template_path = template_metadata.file_path

        document = Document(template_path)
        
        def set_font_size(doc, size):
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size)
                    
                    
        # Insert the required text at the beginning of the document
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Протокол МРТ-Исследования № {reportnumber}")
        run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add patient information
        p = document.add_paragraph()
        p.add_run("ФИО пациента: ").bold = True
        p.add_run(name)
        p.paragraph_format.space_before = Pt(0)  
        p.paragraph_format.space_after = Pt(0)

        p = document.add_paragraph()
        p.add_run("Дата рождения: ").bold = True
        p.add_run(modified_birthdate)
        p.add_run(" г.р.")
        p.paragraph_format.space_before = Pt(0)  
        p.paragraph_format.space_after = Pt(0)
        
        p = document.add_paragraph()
        p.add_run("Вид исследования: ").bold = True
        p.add_run(report_type)
        
        # Добавляем сторону, если она указана
        if side == "right":
            p.add_run(" правого")
        elif side == "left":
            p.add_run(" левого")
            
        p.add_run(" ")
        p.add_run(subtype)
        p.paragraph_format.space_before = Pt(0)  
        p.paragraph_format.space_after = Pt(0)
        
        p = document.add_paragraph()
        p.add_run("Техника сканирования: ").bold = True
        p.add_run(scanParam)
        p.paragraph_format.space_before = Pt(0)  
        p.paragraph_format.space_after = Pt(0)
        
        document.add_paragraph(text)

        # Create a table with one row and three columns
        table = document.add_table(rows=1, cols=3)
        table.autofit = False
        
        # Set column widths
        widths = [Inches(3), Inches(2), Inches(2)]
        for col_idx, width in enumerate(widths):
            for cell in table.columns[col_idx].cells:
                cell.width = width
                
        # First cell: Doctor's name
        cell1 = table.cell(0, 0)
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell1.paragraphs[0]
        run = paragraph.add_run("Врач-рентгенолог: Королёв Д.Г. ")
        run.bold = True
        run.font.size = Pt(12)

        # Second cell: Image
        cell2 = table.cell(0, 1)
        paragraph = cell2.paragraphs[0]
        run = paragraph.add_run()
        if signatura_path:
            run.add_picture(signatura_path, width=Pt(70))

        cell3 = table.cell(0, 2)
        cell3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell3.paragraphs[0]
        run = paragraph.add_run("Дата:  ")
        run.bold = True
        run.font.size = Pt(12)
        run = paragraph.add_run(modified_date_str)
        run.font.size = Pt(12)

        # Set font size for the rest of the document
        set_font_size(document, 12)
        

        # Генерация имени файла
        filename = f"{name}_{report_type}_{subtype}"
        

        # Сохраняем документ во временный файл
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            temp_file_path = tmp_file.name
            document.save(temp_file_path)  # Сохраняем файл в эту временную директорию

        
        folder_name = "reports_word"
          
        result, saved_file_path = file_uploader(temp_file_path, "doc", folder_name, filename, user_id=user_id, profile_id=None, user_email=None)
        
        # Удаляем временный файл после загрузки
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
            
        if "successfully" in result:
            return saved_file_path
        else:
            raise Exception(result)
    except Exception as e:
        logger.error(f"Error in save_to_word: {e}")
        raise Exception(f"Error in save_to_word: {e}")


# Генерация JSON-файлов с уникальными заключениями для всех модальностей, функция 
# вызывается при загрузке сессии и создает файлы и загружает их 
# в OpenAI, кроме того вычищает отовсюду старые файлы старые файлы
def prepare_impression_snippets(profile_id, user_id, user_email, except_words):
    """
    Генерирует новые impression JSON файлы для всех модальностей, 
    удаляет старые файлы (локально, из БД и из OpenAI),
    загружает новые файлы в OpenAI и сохраняет file_ids в session.
    """
    logger.info(f"(функция prepare_impression_snippets) 🚀 Запущена подготовка impression snippets для профиля {profile_id}")
    from tasks.extensions import celery
    client = OpenAI(api_key=celery.conf.OPENAI_API_KEY)
    modalities = ["CT", "MRI", "XRAY"]
    
    if not user_id:
        user_id = current_user.id if current_user.is_authenticated else None
        if not user_id:
            logger.error(f"(функция prepare_impression_snippets) ❌ Не удалось определить user_id для подготовки impression snippets.")
            raise Exception("User ID is required for preparing impression snippets.")
    
    for modality in modalities:
        file_key = f"user:{user_id}:impression_file_id:{modality}"
        
        try:
            logger.info(f"(функция prepare_impression_snippets) 🔄 Работаем с модальностью: {modality}")

            # 1. Удаляем старые файлы
            old_files = FileMetadata.query.filter_by(
                profile_id=profile_id,
                file_description=f"impressions_snippets_{modality}"
            ).all()

            for old_file in old_files:
                if old_file.ai_file_id:
                    try:
                        client.files.delete(old_file.ai_file_id)
                        logger.info(f"(функция prepare_impression_snippets) 🗑 Удалён файл из OpenAI: {old_file.ai_file_id}")
                    except Exception as e:
                        logger.warning(f"(функция prepare_impression_snippets) ⚠️ Не удалось удалить файл {old_file.ai_file_id} из OpenAI: {e}")
                if os.path.exists(old_file.file_path):
                    os.remove(old_file.file_path)
                    logger.info(f"(функция prepare_impression_snippets) 🗑 Удалён локальный файл: {old_file.file_path}")
                    # --- Удаляем папку, если она пуста ---
                    folder_path = os.path.dirname(old_file.file_path)
                    try:
                        # Если папка пуста — удаляем
                        if os.path.isdir(folder_path) and not os.listdir(folder_path):
                            os.rmdir(folder_path)
                            logger.info(f"(функция prepare_impression_snippets) 🗑 Удалена пустая папка: {folder_path}")
                    except Exception as e:
                        logger.warning(f"(функция prepare_impression_snippets) ⚠️ Не удалось удалить папку {folder_path}: {e}")
                db.session.delete(old_file)
            db.session.commit()
            
            # Удаляем старый ключ из Redis, только поле удаления локального файла и файла в OpenAI
            try:
                redis_delete(file_key)
            except Exception as e:
                logger.warning(f"(функция prepare_impression_snippets) ⚠️ Ошибка при удалении ключа {file_key} из Redis: {e}")
            logger.info(f"(функция prepare_impression_snippets) 🗑 Удалены старые файлы для {modality}")

            # 2. Генерируем новый JSON файл
            try:
                new_file_path = generate_impression_json(modality, profile_id, user_id, user_email, except_words)
            except Exception as e:
                logger.error(f"(функция prepare_impression_snippets) ❌ Ошибка при генерации JSON для {modality}: {e}")
                continue

            # 3. Загружаем в OpenAI
            with open(new_file_path, "rb") as f:
                upload = client.files.create(file=f, purpose="assistants")
                new_file_id = upload.id

            # 4. Сохраняем в FileMetadata + session
            new_file_meta = FileMetadata.query.filter_by(
                profile_id=profile_id,
                file_path=new_file_path
            ).first()
            if new_file_meta:
                new_file_meta.ai_file_id = new_file_id
                db.session.commit()
            try:
                redis_set(file_key, new_file_id, ex=60*60*24*30)  # Сохраняем в Redis на 30 дней
                logger.info(f"(функция prepare_impression_snippets) ✅ Загружен и сохранён файл для {modality}: {new_file_id}")
            except Exception as e:
                logger.error(f"(функция prepare_impression_snippets) ❌ Ошибка при сохранении файла в Redis для {modality}: {e}")
                continue

        except Exception as e:
            logger.error(f"(функция prepare_impression_snippets) ❌ Ошибка при обработке модальности {modality}: {e}")
            db.session.rollback()

    logger.info(f"(функция prepare_impression_snippets) 🎉 Подготовка impression snippets завершена для профиля {profile_id}")



# Генерация JSON-файла с уникальными заключениями (использую для индивидуализации ИИ)
def generate_impression_json(modality, profile_id, user_id, user_email, except_words):
    """
    Сбор уникальных заключений (head/body/tail) для заданной модальности и профиля.
    Сохраняет результат в JSON-файл и возвращает путь к нему.

    Args:
        profile_id (int): ID профиля пользователя.
        modality (str): Тип исследования (например, "CT", "MRI", "XRAY").
    
    Returns:
    
        str: Путь к JSON-файлу с уникальными заключениями.
    """
    logger.info(f"(функция generate_impression_json) 🚀 Запущена генерация impression JSON для профиля {profile_id} и модальности {modality}")
    if not profile_id:
        logger.error("(функция generate_impression_json) ❌ Не указан profile_id для генерации impression JSON.")
        return None
    similarity_threshold = 95
    unique_sentences = set()
    cleaned_sentences = []
    
    report_type_id = {"CT": "15", "MRI": "14", "XRAY": "18"}.get(modality.upper(), None)
    
    if not report_type_id:
        logger.error(f"(функция generate_impression_json) ❌ Неизвестная модальность: {modality}. Допустимые значения: CT, MRI, XRAY.")
        raise Exception(f"Unknown modality: {modality}")
        
    subtypes = ReportSubtype.find_by_report_type(report_type_id)
    if not subtypes:
        logger.error(f"(функция generate_impression_json) ❌ Не найдены подтипы для типа отчёта {report_type_id}.")
        raise Exception(f"No subtypes found for report type {report_type_id}.")
    subtype_ids = [subtype.id for subtype in subtypes]
        
    reports = Report.query.filter(
        Report.profile_id == profile_id,
        Report.report_subtype.in_(subtype_ids)
    ).all()
    
    if not reports:
        logger.warning(f"(функция generate_impression_json) ❌ No reports found for profile ID {profile_id}.")
        pass

    for report in reports:
        key_words_for_report = KeyWord.get_keywords_for_report(profile_id, report.id)
        key_words = [keyword.key_word for keyword in key_words_for_report]

        for paragraph in report.report_to_paragraphs:
            if not paragraph.is_impression:
                continue

            head_group = paragraph.head_sentence_group
            tail_group = paragraph.tail_sentence_group

            if head_group:
                for head in HeadSentenceGroup.get_group_sentences(head_group.id):
                    _add_if_unique(head["sentence"], key_words, except_words, cleaned_sentences, unique_sentences, similarity_threshold)

            if tail_group:
                for tail in TailSentenceGroup.get_group_sentences(tail_group.id):
                    _add_if_unique(tail["sentence"], key_words, except_words, cleaned_sentences, unique_sentences, similarity_threshold)

            # body из head-группы
            if head_group:
                for head in HeadSentenceGroup.get_group_sentences(head_group.id):
                    body_group_id = head.get("body_sentence_group_id")
                    if body_group_id:
                        for body in BodySentenceGroup.get_group_sentences(body_group_id):
                            _add_if_unique(body["sentence"], key_words, except_words, cleaned_sentences, unique_sentences, similarity_threshold)

    # Финальный JSON
    data = {
        "profile_id": profile_id,
        "modality": modality,
        "impressions": sorted(list(unique_sentences))
    }

    # Сохраняем файл
    import json
    import tempfile
    with tempfile.NamedTemporaryFile("w+", suffix=".json", delete=False, encoding="utf-8") as tmp_file:
        json.dump(data, tmp_file, indent=2, ensure_ascii=False)
        tmp_file_path = tmp_file.name

    # Загружаем через file_uploader
    result, saved_path = file_uploader(
        tmp_file_path,
        file_type="json", 
        folder_name="impression_snippets",
        file_name=f"{modality.lower()}_impressions_user_{profile_id}",
        file_description=f"impressions_snippets_{modality}",
        user_id=user_id, 
        profile_id=profile_id, 
        user_email=user_email
    )
    if "successfully" not in result:
        logger.error(f"(функция generate_impression_json) ❌ Ошибка при загрузке файла: {result}")
        raise Exception(f"Error uploading file: {result}")
    logger.info(f"(функция generate_impression_json) 📂 Файл с уникальными заключениями сохранён: {saved_path}")
    return saved_path




# Функция для извлечения текста из загруженного файла с помощью OCR
def extract_text_from_uploaded_file(file):
    logger.info(f"(extract_text_from_uploaded_file) Начинаем извлечение текста из файла: {file.filename}")
    # Проверка типа файла (разрешён только jpeg/png)
    allowed_ext = {'jpg', 'jpeg', 'png', 'pdf'}
    filename = secure_filename(file.filename)
    ext = filename.rsplit('.', 1)[-1].lower()
    if ext not in allowed_ext:
        logger.error(f"(extract_text_from_uploaded_file) Ошибка: неподдерживаемый тип файла: {ext}")
        return None, f"Unsupported file type: {ext}"

    # Для PDF — сразу возвращаем "не поддерживается"
    if ext == "pdf":
        return None, "PDF files are not supported for OCR yet. Please upload JPG or PNG images."

    # Сохраняем файл во временный файл
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        file.save(tmp.name)
        tmp_path = tmp.name

    text = ""
    try:
        reader = easyocr.Reader(['ru', 'en'])
        result = reader.readtext(tmp_path, detail=0)
        text = "\n".join(result)
    except Exception as e:
        os.unlink(tmp_path)
        logger.error(f"(extract_text_from_uploaded_file) Ошибка при извлечении текста из файла: {str(e)}")
        return None, f"Ошибка извлечения текста из файла: {str(e)}"
    os.unlink(tmp_path)
    logger.info(f"(extract_text_from_uploaded_file) ✅ Текст успешно извлечен.")  # Логируем первые 100 символов текста
    return text, None